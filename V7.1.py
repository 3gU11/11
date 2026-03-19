import streamlit as st
import pandas as pd
import os
import sys
import ast
from datetime import datetime, timedelta
import time
import uuid
import json
import re
import random
import hashlib
import shutil
import mimetypes
import requests
import tempfile
import subprocess
import cv2
import pymysql
from sqlalchemy import create_engine
from sqlalchemy import create_engine, text
from sqlalchemy.exc import IntegrityError, OperationalError

# OCR Dependencies Check
try:
    from paddleocr import PaddleOCR
    import pdfplumber
    import docx
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("OCR dependencies missing. Please install paddleocr, pdfplumber, python-docx, opencv-python.")

# Mammoth for Docx Preview
try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False
    print("Mammoth library missing. Docx preview will fall back to text. (pip install mammoth)")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# 禁用 PaddleOCR 联网模型检查，解决启动慢的问题
os.environ['PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK'] = 'True'

# 尝试导入 plotly
try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_AVAILABLE = True
except (ImportError, ModuleNotFoundError):
    px = None
    go = None
    PLOTLY_AVAILABLE = False

# ==================== 1. 页面配置与 CSS ====================

st.set_page_config(page_title="成品整机管理系统 V7.0 Pro", layout="wide", page_icon="🏭")

# --- 常量配置 ---
# MySQL Configuration (Environment Variables or Default)
MYSQL_HOST     = os.environ.get('MYSQL_HOST',     'localhost')
MYSQL_PORT     = int(os.environ.get('MYSQL_PORT', '3306'))
MYSQL_USER     = os.environ.get('MYSQL_USER',     'root')
MYSQL_PASSWORD = os.environ.get('MYSQL_PASSWORD', '030705')
MYSQL_DB       = os.environ.get('MYSQL_DB',       'rjt3')

# File Paths (Only for temporary or non-DB files)
ARCHIVE_DIR = "shipping_history" # Only for file archiving if needed, but data goes to DB
CONTRACT_DIR = "data/contracts"
MACHINE_ARCHIVE_DIR = "machine_archives"

@st.cache_resource
def get_engine():
    """返回 SQLAlchemy Engine（全局缓存，整个 Streamlit session 复用）"""
    url = (
        f"mysql+pymysql://{MYSQL_USER}:{MYSQL_PASSWORD}"
        f"@{MYSQL_HOST}:{MYSQL_PORT}/{MYSQL_DB}"
        f"?charset=utf8mb4"
    )
    return create_engine(url, pool_pre_ping=True, pool_recycle=3600)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CONTRACT_ABS_DIR = os.path.join(BASE_DIR, CONTRACT_DIR)
MACHINE_ARCHIVE_ABS_DIR = os.path.join(BASE_DIR, MACHINE_ARCHIVE_DIR)

if not os.path.exists(CONTRACT_ABS_DIR):
    os.makedirs(CONTRACT_ABS_DIR)

if not os.path.exists(MACHINE_ARCHIVE_ABS_DIR):
    os.makedirs(MACHINE_ARCHIVE_ABS_DIR)

# ==================== MySQL 建表初始化 ====================
# ==================== [用户自定义] 账号与权限字典 ====================
DEFAULT_USERS = {
    "boss": {"password": "888", "role": "Boss", "name": "老板"},
    "admin": {"password": "888", "role": "Admin", "name": "系统管理员"},
    "sales": {"password": "123", "role": "Sales", "name": "销售员"},
    "prod": {"password": "123", "role": "Prod", "name": "仓管/生产"},
}

DEFAULT_ROLE_PERMISSIONS = {
    "Boss":  ["PLANNING", "CONTRACT", "QUERY", "ARCHIVE"],
    "Sales": ["PLANNING", "CONTRACT", "SALES_CREATE", "SALES_ALLOC", "INBOUND", "QUERY"],
    "Prod":  ["INBOUND", "SHIP_CONFIRM", "QUERY", "MACHINE_EDIT", "ARCHIVE"],
}

def init_mysql_tables():
    """首次运行时建表，并写入默认用户（幂等操作，可安全重复调用）"""
    engine = get_engine()
    ddl_statements = [
        """
        CREATE TABLE IF NOT EXISTS finished_goods_data (
            `流水号`        VARCHAR(100) NOT NULL,
            `批次号`        VARCHAR(100) DEFAULT '',
            `机型`          VARCHAR(100) DEFAULT '',
            `状态`          VARCHAR(50)  DEFAULT '',
            `预计入库时间`  VARCHAR(50)  DEFAULT '',
            `更新时间`      VARCHAR(50)  DEFAULT '',
            `占用订单号`    VARCHAR(100) DEFAULT '',
            `客户`          VARCHAR(200) DEFAULT '',
            `代理商`        VARCHAR(200) DEFAULT '',
            `订单备注`      TEXT,
            `机台备注/配置` TEXT,
            `Location_Code` VARCHAR(100) DEFAULT '',
            PRIMARY KEY (`流水号`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS sales_orders (
            `订单号`        VARCHAR(100) NOT NULL,
            `客户名`        VARCHAR(200) DEFAULT '',
            `代理商`        VARCHAR(200) DEFAULT '',
            `需求机型`      TEXT,
            `需求数量`      VARCHAR(20)  DEFAULT '',
            `下单时间`      VARCHAR(50)  DEFAULT '',
            `备注`          TEXT,
            `包装选项`      VARCHAR(100) DEFAULT '',
            `发货时间`      VARCHAR(50)  DEFAULT '',
            `指定批次/来源` VARCHAR(200) DEFAULT '',
            `status`        VARCHAR(50)  DEFAULT 'active',
            `delete_reason` TEXT,
            PRIMARY KEY (`订单号`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS factory_plan (
            `id`            INT NOT NULL AUTO_INCREMENT,
            `合同号`        VARCHAR(100) DEFAULT '',
            `机型`          VARCHAR(100) DEFAULT '',
            `排产数量`      VARCHAR(20)  DEFAULT '',
            `要求交期`      VARCHAR(50)  DEFAULT '',
            `状态`          VARCHAR(50)  DEFAULT '',
            `备注`          TEXT,
            `客户名`        VARCHAR(200) DEFAULT '',
            `代理商`        VARCHAR(200) DEFAULT '',
            `指定批次/来源` VARCHAR(200) DEFAULT '',
            `订单号`        VARCHAR(100) DEFAULT '',
            PRIMARY KEY (`id`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS transaction_log (
            `id`        INT NOT NULL AUTO_INCREMENT,
            `时间`      VARCHAR(50)  DEFAULT '',
            `操作类型`  VARCHAR(200) DEFAULT '',
            `流水号`    VARCHAR(100) DEFAULT '',
            `操作员`    VARCHAR(100) DEFAULT '',
            PRIMARY KEY (`id`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS planning_records (
            `id`         INT NOT NULL AUTO_INCREMENT,
            `order_id`   VARCHAR(100) DEFAULT '',
            `model`      VARCHAR(100) DEFAULT '',
            `plan_info`  TEXT,
            `updated_at` VARCHAR(50)  DEFAULT '',
            PRIMARY KEY (`id`),
            UNIQUE KEY `uq_order_model` (`order_id`, `model`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS contract_records (
            `id`          INT NOT NULL AUTO_INCREMENT,
            `contract_id` VARCHAR(100) DEFAULT '',
            `customer`    VARCHAR(200) DEFAULT '',
            `file_name`   VARCHAR(500) DEFAULT '',
            `file_path`   VARCHAR(1000) DEFAULT '',
            `file_hash`   VARCHAR(64)  DEFAULT '',
            `uploader`    VARCHAR(100) DEFAULT '',
            `upload_time` VARCHAR(50)  DEFAULT '',
            PRIMARY KEY (`id`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS audit_log (
            `id`        INT NOT NULL AUTO_INCREMENT,
            `timestamp` VARCHAR(50)  DEFAULT '',
            `user`      VARCHAR(100) DEFAULT '',
            `ip`        VARCHAR(100) DEFAULT '',
            `action`    VARCHAR(200) DEFAULT '',
            `details`   TEXT,
            PRIMARY KEY (`id`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS users (
            `username`      VARCHAR(100) NOT NULL,
            `password`      VARCHAR(200) DEFAULT '',
            `role`          VARCHAR(50)  DEFAULT '',
            `name`          VARCHAR(100) DEFAULT '',
            `status`        VARCHAR(50)  DEFAULT 'pending',
            `register_time` VARCHAR(50)  DEFAULT '',
            `audit_time`    VARCHAR(50)  DEFAULT '',
            `auditor`       VARCHAR(100) DEFAULT '',
            PRIMARY KEY (`username`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS role_permissions (
            `id`        INT NOT NULL AUTO_INCREMENT,
            `role_id`   VARCHAR(50)  DEFAULT '',
            `func_code` VARCHAR(100) DEFAULT '',
            PRIMARY KEY (`id`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS shipping_history (
            `id`            INT NOT NULL AUTO_INCREMENT,
            `批次号`        VARCHAR(100) DEFAULT '',
            `机型`          VARCHAR(100) DEFAULT '',
            `流水号`        VARCHAR(100) DEFAULT '',
            `状态`          VARCHAR(50)  DEFAULT '',
            `预计入库时间`  VARCHAR(50)  DEFAULT '',
            `更新时间`      VARCHAR(50)  DEFAULT '',
            `占用订单号`    VARCHAR(100) DEFAULT '',
            `客户`          VARCHAR(200) DEFAULT '',
            `代理商`        VARCHAR(200) DEFAULT '',
            `订单备注`      TEXT,
            `机台备注/配置` TEXT,
            `archive_month` VARCHAR(20)  DEFAULT '',
            PRIMARY KEY (`id`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS plan_import (
            `流水号`        VARCHAR(100) NOT NULL,
            `批次号`        VARCHAR(100) DEFAULT '',
            `机型`          VARCHAR(100) DEFAULT '',
            `状态`          VARCHAR(50)  DEFAULT '待入库',
            `预计入库时间`  VARCHAR(50)  DEFAULT '',
            `机台备注/配置` TEXT,
            PRIMARY KEY (`流水号`)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """
    ]
    with engine.begin() as conn:
        for ddl in ddl_statements:
            conn.execute(text(ddl))
        result = conn.execute(text("SHOW COLUMNS FROM finished_goods_data LIKE 'Location_Code'"))
        if result.fetchone() is None:
            conn.execute(text("ALTER TABLE finished_goods_data ADD COLUMN `Location_Code` VARCHAR(100) DEFAULT ''"))
        # 写入默认用户（仅当 users 表为空时）
        result = conn.execute(text("SELECT COUNT(*) FROM users"))
        if result.fetchone()[0] == 0:
            current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            for uid, info in DEFAULT_USERS.items():
                conn.execute(text(
                    "INSERT IGNORE INTO users "
                    "(username, password, role, name, status, register_time, audit_time, auditor) "
                    "VALUES (:u, :p, :r, :n, 'active', :t, :t, 'System')"
                ), {"u": uid, "p": info["password"], "r": info["role"], "n": info["name"], "t": current_time})

        result = conn.execute(text("SELECT COUNT(*) FROM role_permissions"))
        if result.fetchone()[0] == 0:
            for role_id, func_codes in DEFAULT_ROLE_PERMISSIONS.items():
                for func_code in func_codes:
                    conn.execute(
                        text("INSERT IGNORE INTO role_permissions (role_id, func_code) VALUES (:r, :f)"),
                        {"r": role_id, "f": func_code},
                    )
        sales_perms = DEFAULT_ROLE_PERMISSIONS.get("Sales", [])
        conn.execute(text("DELETE FROM role_permissions WHERE role_id='Sales'"))
        for func_code in sales_perms:
            conn.execute(
                text("INSERT IGNORE INTO role_permissions (role_id, func_code) VALUES (:r, :f)"),
                {"r": "Sales", "f": func_code},
            )

# 启动时初始化数据库表
if 'db_initialized' not in st.session_state:
    try:
        init_mysql_tables()
        st.session_state.db_initialized = True
    except Exception as _init_err:
        st.error(f"❌ 数据库初始化失败，请检查 MySQL 连接配置：{_init_err}")
        st.stop()

# 预设比例配置
PRESET_RATIOS = {
    "300C": (["FH-300C"], ["FH-300C", "FR-400G", "FR-500G", "FR-600G"]),
    "400G": (["FR-400G"], ["FH-300C", "FR-400G", "FR-500G", "FR-600G"]),
    "500G": (["FR-500G"], ["FH-300C", "FR-400G", "FR-500G", "FR-600G"]),
    "600G": (["FR-600G"], ["FH-300C", "FR-400G", "FR-500G", "FR-600G"]),
    "400XS": (["FR-400XS(PRO)"], ["FR-400XS(PRO)", "FR-500XS(PRO)", "FR-600XS(PRO)", "FR-7055XS(PRO)", "FR-8055XS(PRO)", "FR-8060XS(PRO)"]),
    "500XS": (["FR-500XS(PRO)"], ["FR-400XS(PRO)", "FR-500XS(PRO)", "FR-600XS(PRO)", "FR-7055XS(PRO)", "FR-8055XS(PRO)", "FR-8060XS(PRO)"]),
    "600XS": (["FR-600XS(PRO)"], ["FR-400XS(PRO)", "FR-500XS(PRO)", "FR-600XS(PRO)", "FR-7055XS(PRO)", "FR-8055XS(PRO)", "FR-8060XS(PRO)"]),
    "大机": (["FR-7055XS(PRO)", "FR-8055XS(PRO)", "FR-8060XS(PRO)"], ["FR-400XS(PRO)", "FR-500XS(PRO)", "FR-600XS(PRO)", "FR-7055XS(PRO)", "FR-8055XS(PRO)", "FR-8060XS(PRO)"])
}

# ==================== [用户自定义] 机型排序规则 ====================
CUSTOM_MODEL_ORDER = [
    # FH Series
    "FH-260C", "FH-300C",
    
    # 400 Series
    "FR-400G", "FR-400XS(PRO)", "FR-400AUTO",
    
    # 500 Series
    "FR-500G", "FR-500XS(PRO)", "FR-500AUTO",
    
    # 600 Series
    "FR-600G", "FR-600XS(PRO)", "FR-600AUTO",
    
    # 7055/8055/8060 Series
    "FR-7055AUTO", "FR-7055XS(PRO)", 
    "FR-8055XS(PRO)","FR-8055AUTO", "FR-8060XS(PRO)",
    
    # Large/Other Series
    "FR-1100XS(PRO)", "FL-1390XS(PRO)", "FL-1610XS", "FR-1080Y"
]

def get_model_rank(model_name):
    # 1. 基础清洗
    clean_name = str(model_name).strip()
    
    # 2. 尝试直接匹配
    if clean_name in CUSTOM_MODEL_ORDER:
        return CUSTOM_MODEL_ORDER.index(clean_name)
    
    # 3. 尝试忽略大小写匹配
    upper_list = [x.upper() for x in CUSTOM_MODEL_ORDER]
    if clean_name.upper() in upper_list:
        return upper_list.index(clean_name.upper())
        
    # 4. 尝试移除空格后匹配 (兼容 "FR-400 XS(PRO)" 这种写法)
    nospace_list = [x.replace(" ", "").upper() for x in CUSTOM_MODEL_ORDER]
    clean_nospace = clean_name.replace(" ", "").upper()
    if clean_nospace in nospace_list:
        return nospace_list.index(clean_nospace)
        
    # 5. 兼容旧版写法 (移除连字符)
    # 比如 CUSTOM_MODEL_ORDER 存的是 "FR-400G"，如果来了 "FR400G"，也让它匹配上
    # 或者反之
    # 既然列表里是带连字符的，我们把 clean_name 加连字符比较难，不如把列表去连字符
    nohyphen_list = [x.replace("-", "").upper() for x in CUSTOM_MODEL_ORDER]
    clean_nohyphen = clean_name.replace("-", "").upper()
    if clean_nohyphen in nohyphen_list:
        return nohyphen_list.index(clean_nohyphen)

    return 9999

# 管理员密码 (保留用于兼容，实际使用 USERS)
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "888")

# ==================== [用户自定义] 账号与权限字典 ====================
# Deprecated: USERS dict is now used only for initialization if CSV is missing


def init_users_csv():
    """兼容入口：MySQL 版本中由 init_mysql_tables() 统一处理，此函数保留以避免调用报错"""
    pass

def get_all_users():
    try:
        with get_engine().connect() as conn:
            df = pd.read_sql("SELECT username, password, role, name, status, register_time, audit_time, auditor FROM users", conn)
        return df.fillna("")
    except (OperationalError, Exception) as e:
        print(f"get_all_users error: {e}")
        return pd.DataFrame(columns=["username", "password", "role", "name", "status", "register_time", "audit_time", "auditor"])

def save_all_users(df):
    """全量覆盖 users 表（DELETE + INSERT）"""
    try:
        engine = get_engine()
        with engine.begin() as conn:
            conn.execute(text("DELETE FROM users"))
            if not df.empty:
                df.fillna("").to_sql('users', conn, if_exists='append', index=False, method='multi')
        return True
    except (OperationalError, Exception) as e:
        print(f"save_all_users error: {e}")
        return False

def register_user(username, password, role, name):
    try:
        with get_engine().connect() as conn:
            result = conn.execute(text("SELECT username FROM users WHERE username=:u"), {"u": username})
            if result.fetchone():
                return False, "用户名已存在"
    except (OperationalError, Exception) as e:
        return False, f"系统错误: {e}"

    new_row = {
        "username": username, "password": password, "role": role, "name": name,
        "status": "pending",
        "register_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "audit_time": "", "auditor": ""
    }
    try:
        pd.DataFrame([new_row]).to_sql('users', get_engine(), if_exists='append', index=False, method='multi')
        return True, "注册成功，请等待管理员审核"
    except (IntegrityError, OperationalError, Exception) as e:
        return False, f"系统错误，保存失败: {e}"

def verify_login(username, password):
    df = get_all_users()
    user = df[df['username'] == username]
    
    if user.empty:
        return False, "用户不存在", None
    
    user_row = user.iloc[0]
    if str(user_row['password']) != str(password):
        return False, "密码错误", None
        
    status = str(user_row['status'])
    if status == 'active':
        return True, "登录成功", user_row
    elif status == 'pending':
        return False, "账户待审核，请联系管理员", None
    elif status == 'rejected':
        return False, "账户审核未通过", None
    else:
        return False, f"账户状态异常: {status}", None

# --- Helper for Archive Preview ---
def render_archive_preview(sn):
    """Render archive photos for a machine SN in an interactive grid."""
    if not sn: return
    archive_path = os.path.join(MACHINE_ARCHIVE_ABS_DIR, sn)
    if not os.path.exists(archive_path):
        st.info(f"🚫 该机台 ({sn}) 暂无照片存档")
        return
    
    # Get images
    try:
        all_files = os.listdir(archive_path)
    except Exception:
        st.error(f"无法读取目录: {archive_path}")
        return
        
    image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.gif']
    images = [f for f in all_files if os.path.splitext(f)[1].lower() in image_extensions]
    # Sort by mtime desc
    images.sort(key=lambda x: os.path.getmtime(os.path.join(archive_path, x)), reverse=True)
    
    count = len(images)
    st.markdown(f"**📸 照片总数: {count} 张**")
    
    if count == 0:
        st.warning("该机台目录下无图片文件")
        return

    # Pagination Logic
    BATCH_SIZE = 8
    show_all_key = f"show_all_photos_{sn}"
    
    # Display toggle for show all
    is_show_all = st.checkbox("显示全部照片", key=show_all_key)
    
    display_images = images if is_show_all else images[:BATCH_SIZE]
    
    cols = st.columns(4)
    for idx, img_name in enumerate(display_images):
        img_path = os.path.join(archive_path, img_name)
        with cols[idx % 4]:
            try:
                st.image(img_path, caption=img_name, use_container_width=True)
            except Exception:
                st.error("加载失败")
    
    if count > BATCH_SIZE and not is_show_all:
        st.caption(f"还有 {count - BATCH_SIZE} 张照片未显示，勾选上方选框查看全部。")

def get_role_permissions(role):
    """获取角色权限列表"""
    if role == "Admin":
        return [
            "PLANNING", "CONTRACT", "QUERY", "ARCHIVE",
            "SALES_CREATE", "INBOUND", "SALES_ALLOC",
            "SHIP_CONFIRM", "MACHINE_EDIT"
        ]
    perms = []
    try:
        with get_engine().connect() as conn:
            df_perm = pd.read_sql(
                text("SELECT func_code FROM role_permissions WHERE role_id=:role"),
                conn, params={"role": role}
            )
        perms = df_perm['func_code'].tolist()
    except (OperationalError, Exception) as e:
        print(f"Error loading permissions: {e}")
    if (not perms) and (role in DEFAULT_ROLE_PERMISSIONS):
        return DEFAULT_ROLE_PERMISSIONS[role]
    return perms

# Function Code to Page/UI Mapping
FUNC_MAP = {
    "PLANNING": {"label": "👑 生产统筹", "page": "boss_planning", "class": "boss-btn"},
    "CONTRACT": {"label": "🏭 合同管理", "page": "production", "class": "production-btn"},
    "QUERY": {"label": "🔍 库存查询", "page": "query", "class": "query-btn"},
    "ARCHIVE": {"label": "📂 机台档案", "page": "machine_archive", "class": "machine-edit-btn"},
    "SALES_CREATE": {"label": "📝 销售下单", "page": "sales_create", "class": "sales-create-btn"},
    "INBOUND": {"label": "📥 成品入库", "page": "inbound", "class": "inbound-btn"},
    "SALES_ALLOC": {"label": "📦 订单配货", "page": "sales_alloc", "class": "sales-alloc-btn"},
    "SHIP_CONFIRM": {"label": "🚛 发货复核", "page": "ship_confirm", "class": "ship-btn"},
    "MACHINE_EDIT": {"label": "🛠️ 机台编辑", "page": "machine_edit", "class": "machine-edit-btn"}
}

if not os.path.exists(ARCHIVE_DIR):
    os.makedirs(ARCHIVE_DIR)

# --- 🎨 CSS 样式 ---
st.markdown("""
    <style>
    .block-container { padding-top: 2rem !important; max-width: 100% !important; }
    html, body, [class*="css"] { font-family: 'Segoe UI', sans-serif; }
    .stTextInput label, .stSelectbox label, .stNumberInput label, .stTextArea label, .stRadio label { font-size: 16px !important; font-weight: 600 !important; }
    
    .big-btn button { height: 100px !important; width: 100% !important; font-size: 20px !important; border-radius: 8px !important; }
    .boss-btn button { border: 2px solid #FFD700 !important; color: #DAA520 !important; background-color: #FFFACD !important;}
    .inbound-btn button { border: 2px solid #4CAF50 !important; color: #4CAF50 !important; }
    .sales-create-btn button { border: 2px solid #673AB7 !important; color: #673AB7 !important; } 
    .sales-alloc-btn button { border: 2px solid #9C27B0 !important; color: #9C27B0 !important; } 
    .ship-btn button { border: 2px solid #E91E63 !important; color: #E91E63 !important; }
    .query-btn button { border: 2px solid #FF9800 !important; color: #FF9800 !important; }
    .production-btn button { border: 2px solid #d32f2f !important; color: #d32f2f !important; }
    .machine-edit-btn button { border: 2px solid #607D8B !important; color: #607D8B !important; }
    
    .order-card { background-color: #f0f2f6; padding: 15px; border-radius: 8px; margin-bottom: 15px; border-left: 5px solid #9C27B0; }
    .boss-plan-card { background-color: #FFF8DC; padding: 15px; border-radius: 8px; margin-bottom: 15px; border-left: 5px solid #FFD700; }
    .sub-alloc-card { background-color: #ffffff; padding: 10px; border-radius: 6px; margin-top: 10px; border: 1px solid #e0e0e0; }
    .urgent-alert { padding: 10px; background-color: #ffebee; color: #c62828; border-radius: 5px; border: 1px solid #ef9a9a; margin-bottom: 10px; font-weight: bold; text-align: center;}
    
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    </style>
""", unsafe_allow_html=True)

# ==================== 2. 数据层逻辑 ====================

def get_planning_records():
    try:
        with get_engine().connect() as conn:
            df = pd.read_sql(
                "SELECT order_id, model, plan_info, updated_at FROM planning_records", conn
            )
        return df.fillna("")
    except (OperationalError, Exception):
        return pd.DataFrame(columns=["order_id", "model", "plan_info", "updated_at"])

def save_planning_record(order_id, model, plan_info):
    """
    Upsert planning record for a specific order and model.
    plan_info should be a string (e.g., JSON representation of allocation).
    """
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        with get_engine().begin() as conn:
            # Check if record exists
            result = conn.execute(
                text("SELECT 1 FROM planning_records WHERE order_id=:oid AND model=:m"),
                {"oid": str(order_id), "m": str(model)}
            )
            if result.fetchone():
                # Update existing
                conn.execute(
                    text("UPDATE planning_records SET plan_info=:pi, updated_at=:ua "
                         "WHERE order_id=:oid AND model=:m"),
                    {"pi": str(plan_info), "ua": current_time,
                     "oid": str(order_id), "m": str(model)}
                )
            else:
                # Insert new
                conn.execute(
                    text("INSERT INTO planning_records (order_id, model, plan_info, updated_at) "
                         "VALUES (:oid, :m, :pi, :ua)"),
                    {"oid": str(order_id), "m": str(model),
                     "pi": str(plan_info), "ua": current_time}
                )
    except (IntegrityError, OperationalError, Exception) as e:
        print(f"Error saving planning record: {e}")

def get_data():
    _COLS = ["批次号", "机型", "流水号", "状态", "预计入库时间", "更新时间",
             "占用订单号", "客户", "代理商", "订单备注", "机台备注/配置", "Location_Code"]
    try:
        with get_engine().connect() as conn:
            df = pd.read_sql("SELECT * FROM finished_goods_data", conn)
        if df.empty:
            return pd.DataFrame(columns=_COLS)
        df = df.fillna("")
        for col in _COLS:
            if col not in df.columns:
                df[col] = ""
        # 兼容旧字段名
        if '备注' in df.columns and '订单备注' not in df.columns:
            df.rename(columns={'备注': '订单备注'}, inplace=True)
        try:
            df = df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)
        except:
            pass
        return df
    except (OperationalError, Exception) as e:
        st.error(f"数据读取失败: {e}")
        return pd.DataFrame(columns=_COLS)

def get_orders():
    _COLS = ["订单号", "客户名", "代理商", "需求机型", "需求数量", "下单时间",
             "备注", "包装选项", "发货时间", "指定批次/来源", "status", "delete_reason"]
    try: 
        with get_engine().connect() as conn:
            df = pd.read_sql("SELECT * FROM sales_orders", conn)
        for col in _COLS:
            if col not in df.columns: df[col] = ""
        
        df = df.fillna("")
        
        mask = (df['status'] == "") | (df['status'].isna())
        if mask.any():
            df.loc[mask, 'status'] = "active"
            
        return df
    except (OperationalError, Exception): return pd.DataFrame(columns=_COLS)

def get_factory_plan():
    _COLS = ["合同号", "机型", "排产数量", "要求交期", "状态", "备注",
             "客户名", "代理商", "指定批次/来源", "订单号"]
    try:
        with get_engine().connect() as conn:
            df = pd.read_sql("SELECT * FROM factory_plan", conn)
        for col in _COLS:
            if col not in df.columns: df[col] = ""
        return df.fillna("").drop(columns=['id'], errors='ignore')
    except (OperationalError, Exception):
        return pd.DataFrame(columns=_COLS)

def save_data(df):
    """全量覆盖 finished_goods_data 表（事务内 DELETE + INSERT）"""
    _COLS = ["批次号", "机型", "流水号", "状态", "预计入库时间", "更新时间",
             "占用订单号", "客户", "代理商", "订单备注", "机台备注/配置", "Location_Code"]
    try:
        df = df.drop_duplicates(subset=['流水号'], keep='last')
        df = df.fillna("")
        for col in _COLS:
            if col not in df.columns: df[col] = ""
        with get_engine().begin() as conn:
            result = conn.execute(text("SHOW COLUMNS FROM finished_goods_data LIKE 'Location_Code'"))
            if result.fetchone() is None:
                conn.execute(text("ALTER TABLE finished_goods_data ADD COLUMN `Location_Code` VARCHAR(100) DEFAULT ''"))
            conn.execute(text("DELETE FROM finished_goods_data"))
            if not df.empty:
                df[_COLS].to_sql('finished_goods_data', conn, if_exists='append',
                                 index=False, method='multi', chunksize=500)
    except (OperationalError, Exception) as e:
        st.error(f"保存失败: {e}")
        raise

def save_orders(df):
    """全量覆盖 sales_orders 表"""
    _COLS = ["订单号", "客户名", "代理商", "需求机型", "需求数量", "下单时间",
             "备注", "包装选项", "发货时间", "指定批次/来源", "status", "delete_reason"]
    try:
        df = df.fillna("")
        for col in _COLS:
            if col not in df.columns: df[col] = ""
        with get_engine().begin() as conn:
            conn.execute(text("DELETE FROM sales_orders"))
            if not df.empty:
                df[_COLS].to_sql('sales_orders', conn, if_exists='append',
                                 index=False, method='multi', chunksize=500)
    except (OperationalError, Exception) as e: st.error(f"订单保存失败: {e}")

def save_factory_plan(df):
    """全量覆盖 factory_plan 表"""
    _COLS = ["合同号", "机型", "排产数量", "要求交期", "状态", "备注",
             "客户名", "代理商", "指定批次/来源", "订单号"]
    try:
        df = df.fillna("")
        for col in _COLS:
            if col not in df.columns: df[col] = ""
        with get_engine().begin() as conn:
            conn.execute(text("DELETE FROM factory_plan"))
            if not df.empty:
                df[_COLS].to_sql('factory_plan', conn, if_exists='append',
                                 index=False, method='multi', chunksize=500)
    except (OperationalError, Exception) as e: st.error(f"排产计划保存失败: {e}")

def append_log(action, sn_list):
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    operator = st.session_state.get("operator_name", "Unknown")
    new_logs = [{"时间": current_time, "操作类型": action, "流水号": sn, "操作员": operator} for sn in sn_list]
    if new_logs:
        try:
            pd.DataFrame(new_logs).to_sql(
                'transaction_log', get_engine(), if_exists='append',
                index=False, method='multi'
            )
        except (OperationalError, Exception) as e:
            print(f"append_log error: {e}")

def render_module_logs(filter_keywords):
    """
    渲染模块专属日志
    :param filter_keywords: list of strings, e.g. ["配货锁定", "自动入库"]
    """
    try:
        with get_engine().connect() as conn:
            df_log = pd.read_sql(
                "SELECT 时间, 操作类型, 流水号, 操作员 FROM transaction_log ORDER BY 时间 DESC LIMIT 500",
                conn
            )
        if df_log.empty: return
        
        # 过滤
        if filter_keywords:
            # 只要包含任意一个关键词即可
            mask = df_log['操作类型'].apply(lambda x: any(k in str(x) for k in filter_keywords))
            df_show = df_log[mask]
        else:
            df_show = df_log
            
        if not df_show.empty:
            with st.expander("📜 近期操作日志 (Recent Logs)", expanded=False):
                st.dataframe(df_show, use_container_width=True, hide_index=True)
    except (OperationalError, Exception) as e:
        st.error(f"日志加载失败: {e}")

def archive_shipped_data(df_shipped):
    """将已发货数据存入 shipping_history 表"""
    try:
        current_month = datetime.now().strftime("%Y_%m")
        df_shipped = df_shipped.copy()
        df_shipped['archive_month'] = current_month
        df_shipped.fillna("").to_sql(
            'shipping_history', get_engine(), if_exists='append',
            index=False, method='multi', chunksize=500
        )
    except (OperationalError, Exception) as e:
        print(f"archive_shipped_data error: {e}")


def get_import_staging():
    try:
        with get_engine().connect() as conn:
            df = pd.read_sql("SELECT * FROM plan_import", conn)
        return df.fillna("")
    except (OperationalError, Exception) as e:
        st.error(f"读取待入库数据失败: {e}")
        return pd.DataFrame(columns=["流水号", "批次号", "机型", "状态", "预计入库时间", "机台备注/配置"])

def save_import_staging(df):
    try:
        with get_engine().begin() as conn:
            conn.execute(text("DELETE FROM plan_import"))
            if not df.empty:
                df.to_sql('plan_import', conn, if_exists='append', index=False, method='multi')
    except (OperationalError, Exception) as e:
        st.error(f"保存待入库数据失败: {e}")

def append_import_staging(df):
    if df is None or df.empty:
        return 0
    try:
        df = df.copy()
        df["流水号"] = df["流水号"].astype(str).str.strip()
        df = df[df["流水号"] != ""]
        if df.empty:
            return 0
        df = df.drop_duplicates(subset=["流水号"], keep="last")
        with get_engine().begin() as conn:
            existing_df = pd.read_sql("SELECT 流水号 FROM plan_import", conn)
            existing_sns = set(existing_df["流水号"].astype(str).str.strip().tolist()) if not existing_df.empty else set()
            df_to_append = df[~df["流水号"].isin(existing_sns)].copy()
            if df_to_append.empty:
                return 0
            df_to_append.to_sql('plan_import', conn, if_exists='append', index=False, method='multi')
            return len(df_to_append)
    except (OperationalError, Exception) as e:
        st.error(f"追加待入库数据失败: {e}")
        raise

def clear_import_staging():
    try:
        with get_engine().begin() as conn:
            conn.execute(text("DELETE FROM plan_import"))
    except (OperationalError, Exception) as e:
        st.error(f"清空待入库数据失败: {e}")

# --- 辅助：需求解析 ---
def parse_requirements(model_str, total_qty_str):
    reqs = {}
    m_str = str(model_str)
    if ":" in m_str: 
        try:
            items = m_str.split(";")
            for item in items:
                if ":" in item:
                    k, v = item.split(":")
                    reqs[k.strip()] = int(v)
        except: 
            reqs = {m_str: int(float(total_qty_str)) if total_qty_str else 0}
    else:
        try: q = int(float(total_qty_str))
        except: q = 0
        reqs[m_str] = q
    return reqs

# --- 下单逻辑 ---
def create_sales_order(customer, agent, model_data, note, pack_option="", delivery_time="", source_batch=""):
    odf = get_orders()
    order_id = f"SO-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:4].upper()}"
    
    final_model_str = ""
    total_qty = 0
    
    if isinstance(model_data, dict):
        parts = []
        for m, q in model_data.items():
            parts.append(f"{m}:{q}")
            total_qty += int(q)
        final_model_str = ";".join(parts)
    else:
        final_model_str = str(model_data)
        pass 

    new_row = {
        "订单号": order_id, "客户名": customer, "代理商": agent, 
        "需求机型": final_model_str, 
        "需求数量": str(total_qty) if isinstance(model_data, dict) else "0",
        "下单时间": datetime.now().strftime("%Y-%m-%d %H:%M"), "备注": note,
        "包装选项": pack_option, "发货时间": delivery_time,
        "指定批次/来源": source_batch
    }
    odf = pd.concat([odf, pd.DataFrame([new_row])], ignore_index=True)
    save_orders(odf)
    return order_id

def allocate_inventory(order_id, customer, agent, selected_sns):
    df = get_data()
    orders = get_orders()
    order_note = ""
    target_order = orders[orders['订单号'] == order_id]
    if not target_order.empty: order_note = str(target_order.iloc[0]['备注'])
    
    current_status_df = df[df['流水号'].isin(selected_sns)]
    pending_inbound_sns = current_status_df[current_status_df['状态'] == '待入库']['流水号'].tolist()
    if pending_inbound_sns: append_log("直接配货-自动入库", pending_inbound_sns)

    mask = df['流水号'].isin(selected_sns)
    df.loc[mask, '状态'] = '待发货'
    df.loc[mask, '占用订单号'] = order_id
    df.loc[mask, '客户'] = customer
    df.loc[mask, '代理商'] = agent
    df.loc[mask, '订单备注'] = order_note
    df.loc[mask, '更新时间'] = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    save_data(df)
    append_log(f"配货锁定-{order_id}", selected_sns)

def revert_to_inbound(selected_sns, reason="撤回操作"):
    df = get_data()
    mask = df['流水号'].isin(selected_sns)
    df.loc[mask, '状态'] = '待入库'
    df.loc[mask, '占用订单号'] = ""
    df.loc[mask, '客户'] = ""
    df.loc[mask, '代理商'] = ""
    df.loc[mask, '订单备注'] = ""
    df.loc[mask, '更新时间'] = datetime.now().strftime("%Y-%m-%d %H:%M")
    save_data(df)
    append_log(f"{reason}-退回待入库", selected_sns)

def update_sales_order(order_id, new_data, force_unbind=False):
    df = get_data()
    mask_alloc = (df['占用订单号'] == order_id) & (df['状态'] != '已出库')
    sns_to_unbind = df.loc[mask_alloc, '流水号'].tolist()
    
    has_allocation = len(sns_to_unbind) > 0
    
    if has_allocation:
        if force_unbind:
            revert_to_inbound(sns_to_unbind, reason=f"订单修改-自动解绑-{order_id}")
        else:
            return False, f"⚠️ 警告：该订单已锁定 {len(sns_to_unbind)} 台库存。修改将导致配货失效，是否继续？"
    
    orders = get_orders()
    idx = orders[orders['订单号'] == order_id].index
    if not idx.empty:
        for col, val in new_data.items():
            if col in orders.columns:
                orders.loc[idx, col] = str(val)
        save_orders(orders)
        msg_extra = f"已解绑 {len(sns_to_unbind)} 台关联机器。" if (has_allocation and force_unbind) else ""
        return True, f"订单更新成功！{msg_extra}"
    return False, "订单未找到"

def process_paste_data(raw_text):
    if not raw_text.strip(): return -1, "内容为空"
    try:
        cleaned_text = raw_text.replace("，", ",")
        lines = cleaned_text.strip().split('\n')
        new_records = []
        for line in lines:
            parts = line.replace('\t', ',').split(',')
            parts = [p.strip() for p in parts if p.strip()]
            if len(parts) >= 3:
                b_id = parts[0] if parts[0] not in ['nan', '', 'NaN'] else "无批次"
                record = { "批次号": b_id, "机型": parts[1], "流水号": parts[2] }
                if len(parts) >= 4: record["状态"] = parts[3]
                if len(parts) >= 5: record["预计入库时间"] = parts[4]
                new_records.append(record)
        
        if not new_records: return -1, "未解析出有效数据"
        df_new = pd.DataFrame(new_records)
        if '状态' not in df_new.columns: df_new['状态'] = '待入库'
        
        save_cols = ["批次号", "机型", "流水号", "状态"]
        if "预计入库时间" in df_new.columns: save_cols.append("预计入库时间")
        
        append_import_staging(df_new[save_cols])
        return 1, f"已解析并添加 {len(new_records)} 条数据到计划表"
    except Exception as e: return -1, f"解析错误: {str(e)}"

def generate_auto_inbound(batch_input, model_input, qty_input, expected_inbound_date, machine_note=""):
    if qty_input <= 0: return -1, "数量必须大于0"
    if not batch_input or not model_input: return -1, "批次号和机型不能为空"
    
    if len(machine_note) > 500: return -1, "机台备注/配置内容过长（最大500字符）"
    machine_note = machine_note.replace("<script>", "").replace("</script>", "")

    month_part = ""
    if "-" in batch_input: month_part = batch_input.split("-")[0]
    else: month_part = batch_input 
    
    target_prefix = f"96-{month_part}-"
    existing_sns = set()
    db_df = get_data()
    if not db_df.empty: existing_sns.update(db_df['流水号'].dropna().tolist())
    
    # Check staging as well
    staging_df = get_import_staging()
    if '流水号' in staging_df.columns:
        existing_sns.update(staging_df['流水号'].dropna().tolist())

    max_seq = 0
    for sn in existing_sns:
        sn = str(sn).strip()
        if sn.startswith(target_prefix):
            try:
                suffix = sn.replace(target_prefix, "")
                seq = int(suffix)
                if seq > max_seq: max_seq = seq
            except: continue
    
    new_records = []
    start_seq = max_seq + 1
    expected_inbound_text = str(expected_inbound_date) if expected_inbound_date else ""
    
    for i in range(qty_input):
        current_seq = start_seq + i
        new_sn = f"{target_prefix}{current_seq}"
        new_records.append({
            "批次号": batch_input, "机型": model_input, "流水号": new_sn,
            "状态": "待入库", "预计入库时间": expected_inbound_text,
            "机台备注/配置": machine_note
        })
        
        # --- Auto Create Archive Folder ---
        sn_folder = os.path.join(MACHINE_ARCHIVE_ABS_DIR, new_sn)
        if not os.path.exists(sn_folder):
            try: os.makedirs(sn_folder, exist_ok=True)
            except: pass
        
    if new_records:
        df_new = pd.DataFrame(new_records)
        append_import_staging(df_new)
        return 1, f"已生成 {qty_input} 条数据 ({new_records[0]['流水号']} ~ {new_records[-1]['流水号']})"
    else: return 0, "生成失败"

# 📋 跟踪单导入模块 — 可插入 V6.py 
# =============================================================================

# ───────────────────────────────────────────────────────────────────────────── 
# 【函数1】解析跟踪单 XLS/XLSX → 标准 DataFrame 
# ───────────────────────────────────────────────────────────────────────────── 
def parse_tracking_xls(uploaded_file) -> tuple[int, str, "pd.DataFrame"]: 
    """ 
    解析瑞钧跟踪单 .xls / .xlsx 文件。 
    返回 (code, message, df) 
      code=1  成功 
      code=-1 失败，df 为空 

    提取列： 
      col0 生产批次（前向填充） 
      col1 机型 
      col2 生产编号（流水号） 
      col3 发货日期 → 写入「机台备注/配置」 
    """ 
    if not OPENPYXL_AVAILABLE:
        return -1, "服务器未安装 openpyxl，无法解析 Excel 文件。", pd.DataFrame()

    suffix = os.path.splitext(uploaded_file.name)[-1].lower() 
    raw_bytes = uploaded_file.read() 

    # ── 统一转为 xlsx ────────────────────────────────────────────────────── 
    if suffix == ".xlsx": 
        tmp_xlsx = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) 
        tmp_xlsx.write(raw_bytes); tmp_xlsx.flush(); tmp_xlsx.close() 
        xlsx_path = tmp_xlsx.name 
    elif suffix == ".xls": 
        tmp_xls = tempfile.NamedTemporaryFile(suffix=".xls", delete=False) 
        tmp_xls.write(raw_bytes); tmp_xls.flush(); tmp_xls.close() 
        out_dir = tempfile.mkdtemp() 
        try: 
            # Check for LibreOffice (soffice) or libreoffice
            result = subprocess.run( 
                ["libreoffice", "--headless", "--convert-to", "xlsx", 
                 tmp_xls.name, "--outdir", out_dir], 
                capture_output=True, timeout=30 
            ) 
            if result.returncode != 0: 
                return -1, f"LibreOffice 转换失败: {result.stderr.decode()}", pd.DataFrame() 
        except FileNotFoundError: 
            return -1, "服务器未安装 LibreOffice，无法转换 .xls 文件。\n请将跟踪单另存为 .xlsx 后重新上传。", pd.DataFrame() 
        finally: 
            try: os.unlink(tmp_xls.name)
            except: pass
        converted = [f for f in os.listdir(out_dir) if f.endswith(".xlsx")] 
        if not converted: 
            return -1, "转换后未找到 xlsx 文件", pd.DataFrame() 
        xlsx_path = os.path.join(out_dir, converted[0]) 
    else: 
        return -1, f"不支持的文件格式: {suffix}，请上传 .xls 或 .xlsx", pd.DataFrame() 

    # ── 解析 xlsx ───────────────────────────────────────────────────────── 
    try: 
        wb = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True) 
        ws = wb.active 
        rows = [] 
        current_batch = "" 
        for r_idx, row in enumerate(ws.iter_rows(values_only=True)): 
            if r_idx == 0:          # 跳过表头行 
                continue 
            batch_val = row[0] if len(row) > 0 else None 
            model_val = row[1] if len(row) > 1 else None 
            sn_val    = row[2] if len(row) > 2 else None 
            note_val  = row[3] if len(row) > 3 else None   # 发货日期 

            if batch_val: 
                current_batch = str(batch_val).strip() 

            if not model_val or not sn_val: 
                continue 

            model = str(model_val).strip() 
            sn    = str(sn_val).strip() 
            note  = str(note_val).strip() if note_val else "" 

            if not sn or sn in ("nan", "None"): 
                continue 

            rows.append({ 
                "批次号":     current_batch, 
                "机型":       model, 
                "流水号":     sn, 
                "机台备注/配置": note, 
            }) 
        wb.close() 
    except Exception as e: 
        return -1, f"解析文件时出错: {e}", pd.DataFrame() 
    finally: 
        try: os.unlink(xlsx_path) 
        except: pass 

    if not rows: 
        return -1, "解析后无有效数据，请检查文件格式", pd.DataFrame() 

    df = pd.DataFrame(rows) 
    return 1, f"解析成功，共 {len(df)} 条记录", df 


# ───────────────────────────────────────────────────────────────────────────── 
# 【函数2】与 finished_goods_data 比对，返回新增条目 DataFrame 
# ───────────────────────────────────────────────────────────────────────────── 
def diff_tracking_vs_inventory(tracking_df: "pd.DataFrame") -> "pd.DataFrame": 
    """ 
    比对跟踪单与成品库，返回流水号不在库中的新条目。 
    新条目默认状态 = '待入库'，预计入库时间为空。 
    """ 
    db_df = get_data()
    staging_df = get_import_staging()
    existing_sns = set(db_df["流水号"].astype(str).str.strip().tolist()) if not db_df.empty else set()
    if not staging_df.empty and "流水号" in staging_df.columns:
        existing_sns.update(staging_df["流水号"].astype(str).str.strip().tolist())

    mask = ~tracking_df["流水号"].astype(str).str.strip().isin(existing_sns) 
    new_df = tracking_df[mask].copy() 
    new_df["状态"] = "待入库" 
    new_df["预计入库时间"] = "" 
    # 将「型号」列名对齐为系统内字段「机型」 
    new_df = new_df.rename(columns={"型号": "机型"}) 
    new_df = new_df.reset_index(drop=True) 
    return new_df


# ─────────────────────────────────────────────────────────────────────────────
# 【函数4】UI 渲染 — 插入入库页面 expander 区域
# ─────────────────────────────────────────────────────────────────────────────
def check_prod_admin_permission():
    if st.session_state.role not in ['Admin', 'Prod']:
        st.error("🚫 权限不足！仅限管理员 (Admin) 或 生产/仓管 (Prod) 角色访问。")
        st.stop()

def build_import_payload(selected_df, selected_date):
    if selected_df is None or selected_df.empty:
        return [], "请至少选择 1 条数据"
    if selected_date is None:
        return [], "请选择预计入库日期"
    date_str = selected_date.strftime("%Y-%m-%d") if hasattr(selected_date, "strftime") else str(selected_date)
    payload = []
    for sn in selected_df["流水号"].astype(str).str.strip().tolist():
        if sn:
            payload.append({"trackNo": sn, "expectInDate": date_str})
    if not payload:
        return [], "所选数据缺少有效流水号"
    return payload, ""

def execute_import_transaction_payload(payload, retry_times=1):
    result = {"success": [], "failed": []}
    if not payload:
        return result

    plan_df = get_import_staging().copy()
    if plan_df.empty:
        result["failed"] = [{"trackNo": str(item.get("trackNo", "")), "reason": "待入库清单为空"} for item in payload]
        return result

    plan_df["流水号"] = plan_df["流水号"].astype(str).str.strip()
    staged_map = {row["流水号"]: row for _, row in plan_df.iterrows()}
    payload_map = {}
    for item in payload:
        track_no = str(item.get("trackNo", "")).strip()
        expect_date = str(item.get("expectInDate", "")).strip()
        if not track_no or not expect_date:
            result["failed"].append({"trackNo": track_no, "reason": "参数无效"})
            continue
        payload_map[track_no] = expect_date
        if track_no in staged_map:
            plan_df.loc[plan_df["流水号"] == track_no, "预计入库时间"] = expect_date

    db_df = get_data().copy()
    existing_sns = set(db_df["流水号"].astype(str).str.strip().tolist()) if not db_df.empty else set()
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M")
    rows_to_add = []
    add_track_nos = []
    for track_no, expect_date in payload_map.items():
        if track_no not in staged_map:
            result["failed"].append({"trackNo": track_no, "reason": "待入库清单不存在该流水号"})
            continue
        if track_no in existing_sns:
            result["failed"].append({"trackNo": track_no, "reason": "流水号已在库存中"})
            continue
        row = staged_map[track_no]
        rows_to_add.append({
            "批次号": row.get("批次号", ""),
            "机型": row.get("机型", ""),
            "流水号": track_no,
            "状态": "待入库",
            "预计入库时间": expect_date,
            "更新时间": current_time,
            "占用订单号": "",
            "客户": "",
            "代理商": "",
            "订单备注": "",
            "机台备注/配置": row.get("机台备注/配置", ""),
            "Location_Code": "",
        })
        add_track_nos.append(track_no)

    if rows_to_add:
        df_add = pd.DataFrame(rows_to_add)
        merged_df = pd.concat([db_df, df_add], ignore_index=True)
        merged_df = merged_df.drop_duplicates(subset=['流水号'], keep='first')
        last_error = None
        for _ in range(retry_times + 1):
            try:
                save_data(merged_df)
                last_error = None
                break
            except Exception as e:
                last_error = str(e)
        if last_error is not None:
            for track_no in add_track_nos:
                result["failed"].append({"trackNo": track_no, "reason": f"写入库存失败: {last_error}"})
        else:
            result["success"] = [{"trackNo": track_no} for track_no in add_track_nos]

    success_sns = {item["trackNo"] for item in result["success"]}
    remaining_plan_df = plan_df[~plan_df["流水号"].isin(success_sns)].copy()
    save_import_staging(remaining_plan_df)
    return result

def should_reset_page_selection(prev_page, current_page):
    return prev_page != current_page

def render_tracking_import_module():
    """
    模块二：跟踪单流水号导入
    - 权限控制：仅 Prod/Admin
    - 功能：上传 -> PLAN_IMPORT -> 编辑 -> 确认 -> 写入库存
    """
    check_prod_admin_permission()
    
    st.markdown("### 📋 跟踪单流水号导入模块")
    
    # --- 1. 上传与解析 ---
    with st.expander("📤 上传新跟踪单 (Upload)", expanded=False):
        uploaded = st.file_uploader("选择跟踪单文件 (.xls / .xlsx)", type=["xls", "xlsx"], key="tracking_mod_uploader")
        if uploaded:
            if st.button("🔍 解析并追加到待入库清单", type="primary"):
                with st.spinner("正在解析..."):
                    code, msg, parsed_df = parse_tracking_xls(uploaded)
                    if code == 1:
                        # Diff check logic
                        diff_df = diff_tracking_vs_inventory(parsed_df)
                        
                        if not diff_df.empty:
                            # Append to DB Staging
                            try:
                                append_import_staging(diff_df)
                                st.success(f"✅ 解析成功！已追加 {len(diff_df)} 条记录到待入库清单。")
                                time.sleep(1); st.rerun()
                            except Exception as e:
                                st.error(f"写入待入库清单失败: {e}")
                        else:
                            st.warning("所有解析到的流水号均已在库存中，无需导入。")
                    else:
                        st.error(msg)

    # --- 2. 待入库数据表格展示与编辑 ---
    st.markdown("#### 📝 待入库数据审核 (DB Staging)")
    
    plan_df = get_import_staging().copy()
    
    if plan_df.empty:
        st.info("待入库清单为空，请先上传跟踪单或手动添加。")
    else:
        st.markdown(
            """
            <style>
            div[data-testid="stDataEditor"] table thead tr th:first-child,
            div[data-testid="stDataEditor"] table tbody tr td:first-child {
                width: 40px !important;
                min-width: 40px !important;
                max-width: 40px !important;
                text-align: center !important;
            }
            </style>
            """,
            unsafe_allow_html=True
        )
        plan_df["流水号"] = plan_df["流水号"].astype(str).str.strip()
        filter_col1, filter_col2, filter_col3, filter_col4 = st.columns([2, 1, 1, 2])
        with filter_col1:
            filter_keyword = st.text_input("筛选关键字", value="", key="plan_import_filter_keyword")
        with filter_col2:
            sort_col = st.selectbox("排序字段", ["流水号", "批次号", "机型", "预计入库时间"], index=0, key="plan_import_sort_col")
        with filter_col3:
            sort_asc = st.checkbox("升序", value=True, key="plan_import_sort_asc")
        with filter_col4:
            page_size = st.selectbox("每页条数", [10, 20, 50, 100], index=1, key="plan_import_page_size")

        work_df = plan_df.copy()
        if filter_keyword:
            mask = (
                work_df["流水号"].astype(str).str.contains(filter_keyword, case=False, na=False) |
                work_df["批次号"].astype(str).str.contains(filter_keyword, case=False, na=False) |
                work_df["机型"].astype(str).str.contains(filter_keyword, case=False, na=False)
            )
            work_df = work_df[mask].copy()

        if not work_df.empty:
            work_df = work_df.sort_values(by=sort_col, ascending=sort_asc, kind="stable")

        total_rows = len(work_df)
        total_pages = max(1, (total_rows + page_size - 1) // page_size)
        page_col1, page_col2, page_col3 = st.columns([1, 1, 4])
        with page_col1:
            page_no = st.number_input("页码", min_value=1, max_value=total_pages, value=1, step=1, key="plan_import_page_no")
        with page_col2:
            st.markdown(f"共 {total_pages} 页")
        page_idx = int(page_no) - 1
        start = page_idx * page_size
        end = start + page_size
        page_df = work_df.iloc[start:end].copy()

        if should_reset_page_selection(st.session_state.get("plan_import_prev_page"), page_idx):
            st.session_state["plan_import_selected_map"] = {}
            st.session_state["plan_import_prev_page"] = page_idx

        selected_map = st.session_state.get("plan_import_selected_map", {})
        page_sns = page_df["流水号"].astype(str).tolist()
        for sn in page_sns:
            selected_map.setdefault(sn, False)
        selected_map = {sn: selected_map.get(sn, False) for sn in page_sns}
        st.session_state["plan_import_selected_map"] = selected_map

        selected_count = sum(1 for v in selected_map.values() if v)
        top_left, top_mid, top_right = st.columns([5, 2, 2])
        with top_mid:
            st.markdown(f"**已选 {selected_count} 条**")
        with top_right:
            all_selected = (len(page_sns) > 0 and selected_count == len(page_sns))
            select_all_key = f"plan_import_select_all_{page_idx}"
            select_all = st.checkbox("全选/取消全选", value=all_selected, key=select_all_key)
            if select_all != all_selected:
                selected_map = {sn: select_all for sn in page_sns}
                st.session_state["plan_import_selected_map"] = selected_map

        editor_df = page_df.copy()
        editor_df.insert(0, "选择", [selected_map.get(sn, False) for sn in page_sns])
        edited_plan = st.data_editor(
            editor_df,
            num_rows="fixed",
            hide_index=True,
            use_container_width=True,
            key=f"plan_import_editor_{page_idx}",
            column_config={
                "选择": st.column_config.CheckboxColumn("选择", width="small"),
                "批次号": st.column_config.TextColumn("批次号"),
                "机型": st.column_config.TextColumn("机型"),
                "流水号": st.column_config.TextColumn("流水号"),
                "预计入库时间": st.column_config.TextColumn("预计入库时间"),
                "机台备注/配置": st.column_config.TextColumn("机台备注/配置", width="large"),
            }
        )

        if "选择" in edited_plan.columns:
            current_map = {}
            for _, row in edited_plan.iterrows():
                current_map[str(row["流水号"]).strip()] = bool(row["选择"])
            st.session_state["plan_import_selected_map"] = current_map
            selected_map = current_map
        
        selected_rows = edited_plan[edited_plan["选择"] == True].copy() if "选择" in edited_plan.columns else pd.DataFrame()
        payload_date_col, confirm_btn_col, save_btn_col, msg_col = st.columns([2, 1.5, 1.5, 3])
        with payload_date_col:
            selected_date = st.date_input(
                "预计入库日期",
                value=None,
                min_value=datetime.now().date(),
                format="YYYY-MM-DD",
                key=f"plan_import_date_{page_idx}",
            )
        can_import = (not selected_rows.empty) and (selected_date is not None)
        with msg_col:
            if selected_rows.empty:
                st.warning("请先勾选至少 1 条数据")
            elif selected_date is None:
                st.warning("请选择预计入库日期")
            else:
                st.success(f"已选 {len(selected_rows)} 条，可执行导入")

        with confirm_btn_col:
            if st.button("🚀 确认导入 (Confirm)", type="primary", disabled=not can_import):
                payload, err = build_import_payload(selected_rows, selected_date)
                if err:
                    st.error(err)
                else:
                    import_result = execute_import_transaction_payload(payload, retry_times=1)
                    success_n = len(import_result["success"])
                    failed_n = len(import_result["failed"])
                    if hasattr(st, "toast"):
                        st.toast(f"成功 {success_n} 条，失败 {failed_n} 条")
                    else:
                        st.success(f"成功 {success_n} 条，失败 {failed_n} 条")
                    if failed_n > 0:
                        st.dataframe(pd.DataFrame(import_result["failed"]), use_container_width=True, hide_index=True)
                    time.sleep(0.5)
                    st.rerun()

        col_btns = [save_btn_col]
        with col_btns[0]:
             if st.button("💾 保存修改 (仅保存)", help="将上述修改保存到待入库清单"):
                 try:
                     save_import_staging(edited_plan.drop(columns=["选择"], errors="ignore"))
                     st.success("已保存修改")
                 except Exception as e:
                     st.error(f"保存失败: {e}")

def render_machine_inbound_module():
    """
    模块一：机台入库 (保留原有逻辑)
    - 权限控制：仅 Prod/Admin
    """
    check_prod_admin_permission()
    
    st.markdown("### 🏭 机台入库模块 (扫描入库)")
    
    # Original logic from line 4290
    c_s1, c_s2 = st.columns([3, 1])
    with c_s1: batch = st.text_input("扫描批次号", value=st.session_state.current_batch, key="machine_scan_batch")
    with c_s2: show_all = st.checkbox("显示全部待入库", value=True, key="machine_show_all")
    
    if batch: st.session_state.current_batch = batch
    
    df = get_data()
    # Filter '待入库'
    data = df[df['状态'] == '待入库'].copy()
    
    if not show_all:
        if batch: data = data[data['批次号'] == batch]
        else: data = pd.DataFrame(columns=data.columns)
        
    if not data.empty:
        st.info(f"待入库清单 ({len(data)} 台)")
        # 按机型排序
        data['__rank'] = data['机型'].apply(get_model_rank)
        data = data.sort_values(by=['__rank', '批次号'], ascending=[True, False])
        
        data.insert(0, "✅", False)
        # Use a key to avoid conflict
        res = st.data_editor(
            data[['✅', '批次号', '机型', '流水号', '机台备注/配置']], 
            hide_index=True, 
            use_container_width=True,
            key="machine_inbound_editor"
        )
        
        sel = res[res['✅'] == True]
        
        if not sel.empty:
            if st.button(f"🚀 确认入库 {len(sel)} 台", type="primary", key="btn_confirm_machine_inbound"):
                # Update status
                sns = sel['流水号'].tolist()
                df.loc[df['流水号'].isin(sns), '状态'] = '库存中'
                df.loc[df['流水号'].isin(sns), '更新时间'] = datetime.now().strftime("%Y-%m-%d %H:%M")
                save_data(df)
                append_log("扫描入库", sns)
                st.success("入库成功！"); time.sleep(1); st.rerun()
    else:
        st.info("当前无待入库数据 (或未扫描到对应批次)")


# ==================== Contract & Audit Logic ====================

class OCRProcessor:
    def __init__(self):
        pass

    @staticmethod
    @st.cache_resource(show_spinner="正在加载AI模型，首次运行可能需要几分钟...")
    def get_ocr_model():
        """
        使用 Streamlit 的缓存机制加载模型，防止每次刷新页面都重新加载
        """
        if not OCR_AVAILABLE: return None
        return PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)

    def process_file(self, uploaded_file):
        """处理上传的文件流"""
        if not OCR_AVAILABLE:
            return None, "OCR 依赖未安装"
            
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        
        # 创建临时文件保存上传的数据（因为PaddleOCR和Docx需要文件路径）
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            temp_file.write(uploaded_file.read())
            temp_path = temp_file.name

        full_text = ""
        try:
            if file_ext == '.docx':
                full_text = self._read_docx(temp_path)
            elif file_ext == '.doc':
                full_text = self._read_doc(temp_path)
            elif file_ext == '.pdf':
                full_text = self._read_pdf(temp_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp']:
                full_text = self._ocr_image(temp_path)
            else:
                return None, "不支持的文件格式"
        except Exception as e:
            return None, str(e)
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass
        
        # 重置文件指针，以便后续保存
        uploaded_file.seek(0)

        # 解析字段
        parsed_data = self._parse_fields(full_text)
        return parsed_data, full_text

    def _read_doc(self, path):
        try:
            import win32com.client
            import pythoncom
        except ImportError:
            return "读取 .doc 文件需要安装 pywin32 库 (pip install pywin32) 且系统需安装 Microsoft Word。"

        pythoncom.CoInitialize()
        word = None
        doc = None
        text_content = ""
        try:
            # 使用 DispatchEx 强制启动新实例，避免影响用户当前打开的 Word
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            
            abs_path = os.path.abspath(path)
            doc = word.Documents.Open(abs_path)
            
            # 读取全文
            text_content = doc.Content.Text
            
        except Exception as e:
            text_content = f"读取 .doc 失败: {str(e)}"
        finally:
            if doc:
                try:
                    doc.Close(False)
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()
            
        return text_content


    def _read_docx(self, path):
        doc = docx.Document(path)
        text = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                text.append(" ".join([cell.text for cell in row.cells]))
        return "\n".join(text)

    def _read_pdf(self, path):
        text_content = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text_content += extracted + "\n"
        # 注意：如果PDF是纯图扫描件，pdfplumber提取为空，这里为了演示简化，
        # 实际生产建议检测文本长度，如果过短则调用 pdf2image 转图后再 OCR
        if len(text_content) < 10: 
            return "[提示] 这是一个扫描版PDF，本基础版本暂仅支持文字版PDF提取。"
        return text_content

    def _ocr_image(self, path):
        ocr = self.get_ocr_model()
        
        # 1. 图像预处理 (增强识别率)
        try:
            image = cv2.imread(path)
            if image is None:
                return f"[错误] 无法读取图片: {path}"
                
            # 灰度化
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # 去噪 (高斯模糊)
            denoised = cv2.GaussianBlur(gray, (5, 5), 0)
            
            # 二值化 (自适应阈值，处理光照不均)
            # thresh = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            # 对于合同文档，通常简单的二值化或者保持原图（PaddleOCR内部有处理）即可
            # 但如果图片质量差，适当的锐化有帮助
            # 这里我们尝试保留原图给Paddle，因为PaddleOCR内部集成了很强的预处理
            # 仅当原图识别失败时，可以考虑预处理。
            # 为了稳妥，我们先传原图路径。如果用户反馈有问题，可能是分辨率过大或过小。
            
            # PaddleOCR 建议输入是 numpy array
            result = ocr.ocr(image, cls=True)
            
        except Exception as e:
            return f"[错误] 图片预处理失败: {str(e)}"

        text_lines = []
        if result and result[0]:
            # 按垂直坐标排序，尽可能还原阅读顺序
            # PaddleOCR result: [[[x,y], [x,y]...], (text, conf)]
            # sort by y coordinate of the first point box[0][1]
            sorted_res = sorted(result[0], key=lambda x: x[0][0][1])
            
            for line in sorted_res:
                txt = line[1][0]
                conf = line[1][1]
                if conf > 0.5: # 过滤低置信度
                    text_lines.append(txt)
        else:
            return "[警告] 未识别到有效文字 (可能是图片模糊或方向不对)"
            
        return "\n".join(text_lines)

    def _parse_fields(self, text):
        """字段提取逻辑：正则 + LLM"""
        # 1. 尝试使用 LLM 进行智能提取
        llm_data = self._extract_with_llm(text)
        if llm_data:
            return llm_data
            
        # 2. 如果 LLM 失败，回退到正则规则提取
        st.warning("AI 提取失败，正在使用规则提取...")
        return self._parse_fields_regex(text)

    def _extract_with_llm(self, text):
        """使用大模型进行结构化提取"""
        api_key = "sk-b8ebeee6dc8b4d7eaae1e2502ddf3ff9"
        base_url = "https://api.deepseek.com/chat/completions"
        
        # 截取文本以避免超过 Context Window (保留前 3000 字通常足够)
        truncated_text = text[:3000]
        
        prompt = f"""
        我将提供一段OCR识别的文本（通常是合同或订单），请从中提取以下字段并以JSON格式返回：
        
        {{
            "customer": "需方/客户名称",
            "agent": "代理商（如果有）",
            "deadline": "交货日期/要求交期（格式：YYYY-MM-DD）",
            "global_note": "合同总备注/其他重要条款",
            "items": [
                {{
                    "model": "机型名称（请去除数量等无关字符）",
                    "qty": 数量（整数）,
                    "is_high": true/false（如果包含'加高'字样则为true）,
                    "note": "单行备注（针对该机型的特殊要求）"
                }}
            ]
        }}
        
        注意事项：
        1. 如果未找到某个字段，请填 null 或空字符串，不要填 "未识别"。
        2. "机型"如果是表格形式，请仔细解析每一行。
        3. "加高"通常出现在机型名称后或备注中。
        4. 请只返回纯 JSON，不要包含 markdown 标记。
        
        文本内容如下：
        {truncated_text}
        """
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        payload = {
            "model": "deepseek-chat", # DeepSeek 官方 API 模型名通常为 deepseek-chat 或 deepseek-reasoner
            "messages": [
                {"role": "system", "content": "你是一个智能文档分析助手，擅长从OCR文本中提取结构化合同数据。只返回纯 JSON 格式的数据。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1,
            "max_tokens": 1000,
            "stream": False
        }
        
        try:
            response = requests.post(base_url, headers=headers, json=payload, timeout=30)
            if response.status_code == 200:
                res_json = response.json()
                content = res_json['choices'][0]['message']['content']
                # 清洗 markdown 代码块标记
                content = content.replace("```json", "").replace("```", "").strip()
                return json.loads(content)
            else:
                print(f"API Error: {response.text}")
                return None
        except Exception as e:
            print(f"LLM Error: {e}")
            return None

    def _parse_fields_regex(self, text):
         # 简单兜底
         return {'需方': '未识别', '机型及数量': '未识别', '地址': '未识别', '交货日期': '未识别'}

def audit_log(action, details):
    user = st.session_state.get("operator_name", "Unknown")
    # Try to get IP (Best effort in Streamlit)
    ip = "Local"
    try:
        from streamlit.web.server.websocket_headers import _get_websocket_headers
        headers = _get_websocket_headers()
        if headers:
            ip = headers.get("X-Forwarded-For", headers.get("Remote-Addr", "Local"))
    except: pass
    
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    new_row = {
        "timestamp": timestamp, "user": user, "ip": ip, 
        "action": action, "details": details
    }
    try:
        pd.DataFrame([new_row]).to_sql('audit_log', get_engine(), if_exists='append', index=False, method='multi')
    except (OperationalError, Exception): pass

def get_contract_files(contract_id=None):
    try:
        # 1. Get from DB
        with get_engine().connect() as conn:
            if contract_id:
                df = pd.read_sql(
                    text("SELECT * FROM contract_records WHERE contract_id=:cid"),
                    conn, params={"cid": str(contract_id)}
                )
            else:
                df = pd.read_sql(text("SELECT * FROM contract_records"), conn)
        
        df = df.fillna("")
        
        # 2. Sync with Disk (Folder-based view)
        if contract_id:
            safe_cid = re.sub(r'[\\/*?:"<>|]', "", str(contract_id)).strip()
            folder_paths = [
                os.path.join(BASE_DIR, "data", safe_cid),
                os.path.join(BASE_DIR, "data", "contracts", safe_cid)
            ]
            folder_paths = [p for p in folder_paths if os.path.exists(p) and os.path.isdir(p)]

            if folder_paths:
                disk_files = []
                existing_files = set(df['file_name'].tolist()) if not df.empty else set()
                existing_paths = set(df['file_path'].tolist()) if not df.empty else set()
                seen_disk_paths = set()

                for folder_path in folder_paths:
                    rel_prefix = os.path.relpath(folder_path, BASE_DIR)
                    for f in os.listdir(folder_path):
                        abs_file = os.path.join(folder_path, f)
                        if os.path.isfile(abs_file):
                            rel_file = os.path.join(rel_prefix, f)
                            if rel_file in seen_disk_paths:
                                continue
                            seen_disk_paths.add(rel_file)
                            if (f not in existing_files) and (rel_file not in existing_paths):
                                disk_files.append({
                                    "contract_id": str(contract_id),
                                    "customer": "Unknown",
                                    "file_name": f,
                                    "file_path": rel_file,
                                    "file_hash": "",
                                    "uploader": "Disk/System",
                                    "upload_time": datetime.fromtimestamp(os.path.getmtime(abs_file)).strftime("%Y-%m-%d %H:%M:%S")
                                })
                
                if disk_files:
                    df_disk = pd.DataFrame(disk_files)
                    df = pd.concat([df, df_disk], ignore_index=True)

        return df
    except (OperationalError, Exception) as e:
        print(f"Error getting contract files: {e}")
        return pd.DataFrame(columns=["contract_id", "customer", "file_name", "file_path", "file_hash", "uploader", "upload_time"])

def get_unlinked_contract_folders(known_contract_ids=None):
    rows = []
    known_ids = set(str(x).strip() for x in (known_contract_ids or []) if str(x).strip())
    try:
        data_root = os.path.join(BASE_DIR, "data")
        if not os.path.exists(data_root) or not os.path.isdir(data_root):
            return pd.DataFrame(columns=["合同号", "客户名", "文件数", "最近更新时间"])

        scan_roots = [data_root, os.path.join(data_root, "contracts")]
        seen_ids = set()
        for root in scan_roots:
            if not os.path.exists(root) or not os.path.isdir(root):
                continue
            for folder in os.listdir(root):
                folder_path = os.path.join(root, folder)
                if not os.path.isdir(folder_path):
                    continue
                cid = str(folder).strip()
                if not cid or cid in seen_ids:
                    continue
                seen_ids.add(cid)
                if cid in known_ids:
                    continue

                files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
                if not files:
                    continue

                latest_ts = max(os.path.getmtime(os.path.join(folder_path, f)) for f in files)
                customer = cid.split("_")[0] if "_" in cid else cid
                rows.append({
                    "合同号": cid,
                    "客户名": customer,
                    "文件数": len(files),
                    "最近更新时间": datetime.fromtimestamp(latest_ts).strftime("%Y-%m-%d %H:%M:%S")
                })

        if not rows:
            return pd.DataFrame(columns=["合同号", "客户名", "文件数", "最近更新时间"])
        df = pd.DataFrame(rows)
        return df.sort_values("最近更新时间", ascending=False).reset_index(drop=True)
    except Exception as e:
        print(f"Error scanning unlinked contract folders: {e}")
        return pd.DataFrame(columns=["合同号", "客户名", "文件数", "最近更新时间"])

def save_contract_file(uploaded_file, customer_name, contract_id, uploader_name, convert_to_docx=True):
    # 1. Check size (50MB)
    if uploaded_file.size > 50 * 1024 * 1024:
        return False, "文件超过 50MB 限制"
    
    # 2. Check type
    fname = uploaded_file.name
    ext = os.path.splitext(fname)[1].lower()
    if ext not in ['.pdf', '.doc', '.docx', '.jpg', '.jpeg']:
        return False, "不支持的文件格式 (仅限 PDF, Word, JPG)"
    
    # 3. Calculate Hash
    file_bytes = uploaded_file.getvalue()
    file_hash = hashlib.sha256(file_bytes).hexdigest()
    
    # 4. Save to Disk
    # Modified: Use 'data/{contract_id}' as the folder structure
    # If contract_id is empty, use safe_cust_pending
    
    safe_cust = re.sub(r'[\\/*?:"<>|]', "", str(customer_name)).strip()
    if not safe_cust: safe_cust = "Unknown"
    
    if contract_id:
        folder_name = re.sub(r'[\\/*?:"<>|]', "", str(contract_id)).strip()
    else:
        timestamp_str = datetime.now().strftime("%Y%m%d%H%M%S")
        folder_name = f"{safe_cust}_{timestamp_str}"

    # Use 'data' as the parent directory for contracts
    rel_dir = os.path.join("data", folder_name)
    abs_dir = os.path.join(BASE_DIR, rel_dir)
    
    if not os.path.exists(abs_dir):
        try:
            os.makedirs(abs_dir, exist_ok=True)
        except Exception as e:
            return False, f"创建目录失败: {e}"
        
    save_path = os.path.join(abs_dir, fname)
    rel_save_path = os.path.join(rel_dir, fname)
    
    try:
        with open(save_path, "wb") as f:
            f.write(file_bytes)
    except Exception as e:
        return False, f"保存文件失败: {e}"

    # --- Auto Convert .doc to .docx (External Script) ---
    final_fname = fname
    final_rel_path = rel_save_path
    
    if ext == '.doc' and convert_to_docx:
        # Script path: finished/合同/合同/convert_doc_to_docx.py
        # Assuming V6.py is in finished/
        script_path = os.path.join(BASE_DIR, "合同", "合同", "convert_doc_to_docx.py")
        
        if os.path.exists(script_path):
            try:
                # Run the script in subprocess
                # python script.py "path/to/doc"
                cmd = [sys.executable, script_path, save_path]
                print(f"Running conversion script: {cmd}")
                
                # Run and wait for completion
                result = subprocess.run(cmd, capture_output=True, text=True)
                
                print(f"STDOUT: {result.stdout}")
                print(f"STDERR: {result.stderr}")

                if result.returncode == 0:
                    # Check if docx exists
                    abs_docx_path = os.path.splitext(save_path)[0] + ".docx"
                    if os.path.exists(abs_docx_path):
                        st.toast("✅ .doc 自动转换成功！")
                        print("Conversion successful via external script.")
                        # Script handles deletion of original doc if configured, 
                        # but our modified script does delete_original=True
                        
                        # Update variables
                        final_fname = os.path.splitext(fname)[0] + ".docx"
                        final_rel_path = os.path.splitext(rel_save_path)[0] + ".docx"
                        
                        # Recalculate hash for the new file
                        with open(abs_docx_path, "rb") as f:
                            file_hash = hashlib.sha256(f.read()).hexdigest()
                    else:
                        err_msg = f"转换脚本运行成功但未生成 .docx 文件。\n输出: {result.stdout}"
                        print(err_msg)
                        st.error(err_msg)
                else:
                    err_msg = f"转换脚本运行失败 (Code {result.returncode}):\n{result.stderr}"
                    print(err_msg)
                    st.error(err_msg)
            except Exception as e:
                err_msg = f"执行转换脚本时发生异常: {e}"
                print(err_msg)
                st.error(err_msg)
        else:
            msg = f"未找到转换脚本: {script_path}"
            print(msg)
            st.warning(msg)
        
    # 5. Record to DB
    new_record = {
        "contract_id": str(contract_id) if contract_id else "",
        "customer": str(customer_name),
        "file_name": final_fname,
        "file_path": final_rel_path,
        "file_hash": file_hash,
        "uploader": uploader_name,
        "upload_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    
    try:
        pd.DataFrame([new_record]).to_sql('contract_records', get_engine(), if_exists='append', index=False, method='multi')
        
        audit_log("Upload Contract", f"Uploaded {final_fname} for {customer_name} (ID: {contract_id})")
        return True, "上传成功"
    except (IntegrityError, OperationalError, Exception) as e:
        return False, f"记录保存失败: {e}"

def delete_contract_file(contract_id, file_name):
    try:
        # Check if record exists
        with get_engine().connect() as conn:
            result = conn.execute(
                text("SELECT file_path FROM contract_records WHERE contract_id=:cid AND file_name=:fn"),
                {"cid": str(contract_id), "fn": str(file_name)}
            )
            row = result.fetchone()
        
        if row:
            # Get file path
            rel_path = row[0]
            abs_path = os.path.join(BASE_DIR, rel_path)
            
            # Delete physical file
            if os.path.exists(abs_path):
                try:
                    os.remove(abs_path)
                except Exception as e:
                    return False, f"物理文件删除失败: {e}"
            
            # Update DB
            with get_engine().begin() as conn:
                conn.execute(
                    text("DELETE FROM contract_records WHERE contract_id=:cid AND file_name=:fn"),
                    {"cid": str(contract_id), "fn": str(file_name)}
                )
            
            audit_log("Delete Contract File", f"Deleted {file_name} from {contract_id}")
            return True, "文件已删除"
        else:
            return False, "文件记录未找到"
    except (IntegrityError, OperationalError, Exception) as e:
        return False, f"删除操作出错: {e}"

def clean_expired_contracts():
    # Retention: 3 years
    retention_days = 365 * 3
    cutoff_date = datetime.now() - timedelta(days=retention_days)
    cutoff_str = cutoff_date.strftime("%Y-%m-%d %H:%M:%S")
    
    try:
        # Get expired files
        with get_engine().connect() as conn:
            df = pd.read_sql(
                text("SELECT file_path FROM contract_records WHERE upload_time < :cut"),
                conn, params={"cut": cutoff_str}
            )
        
        if df.empty: return
        
        count = 0
        for _, row in df.iterrows():
            f_path = row['file_path']
            abs_path = os.path.join(BASE_DIR, f_path)
            if os.path.exists(abs_path):
                try:
                    os.remove(abs_path)
                    count += 1
                except: pass
        
        # Delete from DB
        with get_engine().begin() as conn:
            conn.execute(
                text("DELETE FROM contract_records WHERE upload_time < :cut"),
                {"cut": cutoff_str}
            )
            
        if count > 0:
            print(f"[System] Cleaned {count} expired contract files.")
            audit_log("System Cleanup", f"Deleted {count} expired files older than {retention_days} days")
                
    except (OperationalError, Exception) as e:
        print(f"Cleanup error: {e}")

# Run cleanup once per session/day (simple check)
if 'last_cleanup' not in st.session_state:
    clean_expired_contracts()
    st.session_state.last_cleanup = datetime.now().strftime("%Y-%m-%d")

# ==================== 3. 辅助逻辑：权限与通知 ====================

# 1. 登录与会话初始化
if 'current_user' not in st.session_state:
    st.session_state.current_user = None
    st.session_state.role = None
    st.session_state.operator_name = ''
    st.session_state.is_admin = False
    st.session_state.page = 'home'
    st.session_state.current_batch = ''
    st.session_state.permissions = [] # Cache for permissions

def login_form():
    """显示登录表单 (支持注册)"""
    # 使用空容器居中显示
    st.markdown("<br><br><br>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("<h2 style='text-align: center;'>🔐 成品管理系统</h2>", unsafe_allow_html=True)
        st.write("")
        
        tab_login, tab_register = st.tabs(["🔑 登录", "📝 注册新账号"])
        
        with tab_login:
            with st.form("login_form"):
                username = st.text_input("账号")
                password = st.text_input("密码", type="password")
                submitted = st.form_submit_button("登录", use_container_width=True)
                
                if submitted:
                    ok, msg, user_row = verify_login(username, password)
                    if ok:
                        st.session_state.current_user = username
                        st.session_state.role = user_row["role"]
                        st.session_state.operator_name = user_row["name"]
                        st.session_state.is_admin = (user_row["role"] == "Admin")
                        
                        # --- Load Permissions ---
                        perms = get_role_permissions(user_row["role"])
                        st.session_state.permissions = perms
                        
                        st.success(f"{msg}！欢迎 {user_row['name']}")
                        time.sleep(0.5)
                        st.rerun()
                    else:
                        st.error(msg)

        with tab_register:
            with st.form("register_form"):
                r_user = st.text_input("设置账号 (用户名)", help="登录用的唯一ID")
                r_pass = st.text_input("设置密码", type="password")
                r_name = st.text_input("您的姓名 (真实姓名)")
                
                # Role mapping for display
                role_display_map = {"销售员": "Sales", "生产/仓管": "Prod", "老板/管理": "Boss"}
                r_role_display = st.selectbox("申请角色", list(role_display_map.keys()))
                r_role = role_display_map[r_role_display]
                
                reg_submitted = st.form_submit_button("提交注册申请", use_container_width=True)
                
                if reg_submitted:
                    if not r_user or not r_pass or not r_name:
                        st.error("请填写完整信息")
                    else:
                        ok, msg = register_user(r_user, r_pass, r_role, r_name)
                        if ok: st.success(msg)
                        else: st.error(msg)

# 如果未登录，只显示登录页并停止后续渲染
if st.session_state.current_user is None:
    login_form()
    st.stop()  # 阻止后续代码执行

def go_page(name): st.session_state.page = name
def go_home(): st.session_state.page = 'home'

def check_access(required_permission):
    """
    检查当前用户是否有特定权限 (Function Code)
    :param required_permission: e.g. "PLANNING", "SALES_CREATE"
    """
    # Admin bypass (Redundant check but safe)
    if st.session_state.role == "Admin": return

    user_perms = st.session_state.get('permissions', [])
    
    # Check if permission exists
    if required_permission not in user_perms:
        st.error(f"🚫 权限不足！需要权限: {required_permission}")
        st.button("⬅️ 返回首页", on_click=go_home)
        st.stop()

def get_urgent_production_count():
    fp = get_factory_plan()
    if fp.empty: return 0
    today = datetime.now().date()
    target_date = today + timedelta(days=14)
    count = 0
    for _, row in fp.iterrows():
        status = str(row['状态']).strip()
        deadline_str = str(row['要求交期']).strip()
        if status == "未下单" and deadline_str:
            try:
                d_date = datetime.strptime(deadline_str, "%Y-%m-%d").date()
                if d_date <= target_date: count += 1
            except: pass
    return count

def render_file_manager(contract_id, customer_name, default_expanded=True, key_suffix=""):
    """
    渲染合同文件管理组件
    :param contract_id: 合同号
    :param customer_name: 客户名 (用于上传)
    :param default_expanded: 是否默认展开
    :param key_suffix: 避免 key 冲突的后缀
    """
    
    # 容器或折叠面板
    container = st.container()
    if not default_expanded:
        container = st.expander("📎 合同附件 (点击展开)", expanded=False)
    
    with container:
        c_files = get_contract_files(contract_id)
        if default_expanded: 
            st.divider()
            st.markdown("#### 📎 合同附件管理")
        
        if not c_files.empty:
            # Default preview: auto-select the first file if none selected
            preview_key = f"preview_target_{contract_id}{key_suffix}"
            if preview_key not in st.session_state:
                st.session_state[preview_key] = c_files.iloc[0]['file_name']

            # File List View
            for _, f_row in c_files.iterrows():
                f_name = f_row['file_name']
                f_path = f_row['file_path']
                abs_path = os.path.join(BASE_DIR, f_path)
                
                # --- Path Recovery Logic for Legacy/Moved Files ---
                if not os.path.exists(abs_path):
                    # Try removing 'contracts/' prefix if present (Migration scenario)
                    if "contracts" in f_path:
                        alt_path = f_path.replace("contracts/", "").replace("contracts\\", "")
                        alt_abs_path = os.path.join(BASE_DIR, alt_path)
                        if os.path.exists(alt_abs_path):
                            f_path = alt_path
                            abs_path = alt_abs_path
                # ---------------------------------------------------

                col_icon, col_name, col_info, col_act = st.columns([0.5, 3, 2, 2.5])
                
                with col_icon:
                    ext = os.path.splitext(f_name)[1].lower()
                    icon = "📄"
                    if ext in ['.jpg', '.jpeg', '.png']: icon = "🖼️"
                    elif ext == '.pdf': icon = "📕"
                    elif ext in ['.doc', '.docx']: icon = "📝"
                    st.write(f"### {icon}")
                    
                with col_name:
                    st.write(f"**{f_name}**")
                
                with col_info:
                    st.caption(f"👤 {f_row['uploader']}\n🕒 {f_row['upload_time']}")
                    
                with col_act:
                    c_a1, c_a2, c_a3 = st.columns(3)
                    with c_a1:
                        # Download
                        if os.path.exists(abs_path):
                            with open(abs_path, "rb") as f:
                                st.download_button("下载", f.read(), file_name=f_name, key=f"dl_{f_row.name}{key_suffix}", help="下载")
                        else:
                            st.error("文件丢失")
                    with c_a2:
                        # Preview Toggle
                        if st.button("预览", key=f"prev_{f_row.name}{key_suffix}", help="预览"):
                            st.session_state[preview_key] = f_name
                            st.rerun()
                    with c_a3:
                        # Delete
                        if st.button("删除", key=f"del_{f_row.name}{key_suffix}", help="删除"):
                            ok, msg = delete_contract_file(contract_id, f_name)
                            if ok:
                                st.success(msg)
                                if st.session_state.get(preview_key) == f_name:
                                    del st.session_state[preview_key]
                                time.sleep(0.5); st.rerun()
                            else: st.error(msg)
            
            # Preview Area
            target_file_name = st.session_state.get(preview_key)
            if target_file_name:
                # Find the record
                target_rec = c_files[c_files['file_name'] == target_file_name]
                if not target_rec.empty:
                    f_path = target_rec.iloc[0]['file_path']
                    abs_path = os.path.join(BASE_DIR, f_path)
                    
                    # --- Path Recovery for Preview ---
                    if not os.path.exists(abs_path):
                         if "contracts" in f_path:
                             alt_path = f_path.replace("contracts/", "").replace("contracts\\", "")
                             alt_abs_path = os.path.join(BASE_DIR, alt_path)
                             if os.path.exists(alt_abs_path):
                                 abs_path = alt_abs_path
                    # ---------------------------------
                    
                    ext = os.path.splitext(target_file_name)[1].lower()
                    
                    st.info(f"正在预览: {target_file_name}")
                    if os.path.exists(abs_path):
                        if ext in ['.jpg', '.jpeg', '.png']:
                            st.image(abs_path, use_container_width=True)
                        elif ext == '.pdf':
                            import base64
                            with open(abs_path, "rb") as f:
                                base64_pdf = base64.b64encode(f.read()).decode('utf-8')
                            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600" type="application/pdf"></iframe>'
                            st.markdown(pdf_display, unsafe_allow_html=True)
                        elif ext == '.docx':
                            if MAMMOTH_AVAILABLE:
                                with open(abs_path, "rb") as docx_file:
                                    result = mammoth.convert_to_html(docx_file)
                                    st.markdown(result.value, unsafe_allow_html=True)
                            else:
                                st.warning("缺少 mammoth 库，无法预览 docx")
                        else:
                            st.info("此格式不支持在线预览，请下载查看")
                    else:
                        st.error("文件不存在")
        else:
            st.info("暂无附件")

        # Upload New
        with st.expander("📤 上传新附件 (支持覆盖)", expanded=False):
            new_files = st.file_uploader("拖拽文件到此处", accept_multiple_files=True, key=f"new_up_{contract_id}{key_suffix}", type=['pdf', 'doc', 'docx', 'jpg', 'jpeg'])
            overwrite = st.checkbox("遇到同名文件时自动覆盖", value=True, key=f"ov_{contract_id}{key_suffix}")
            
            # 检测 .doc 并显示转换选项
            convert_option = False
            has_doc = False
            if new_files:
                for nf in new_files:
                    if nf.name.lower().endswith('.doc'):
                        has_doc = True
                        break
            
            if has_doc:
                st.warning("⚠️ 检测到旧版 Word (.doc) 格式")
                convert_option = st.checkbox("🔄 自动转换为 .docx (推荐，方便在线预览)", value=True, key=f"cv_{contract_id}{key_suffix}", help="如果不勾选，将保留原始 .doc 格式，可能无法在线预览")

            if new_files and st.button("开始上传", key=f"btn_up_{contract_id}{key_suffix}"):
                cnt = 0
                for nf in new_files:
                    # Check exist
                    if not c_files.empty and nf.name in c_files['file_name'].values:
                        if overwrite:
                            delete_contract_file(contract_id, nf.name)
                        else:
                            st.warning(f"跳过同名文件: {nf.name}"); continue
                    
                    # Pass convert_option
                    ok, msg = save_contract_file(nf, customer_name, contract_id, st.session_state.operator_name, convert_to_docx=convert_option)
                    if ok: cnt += 1
                
                if cnt > 0:
                    st.success(f"成功上传 {cnt} 个文件！")
                    time.sleep(1); st.rerun()

# ==================== 4. 页面路由 ====================

with st.sidebar:
    st.title("🏭 管理系统 V7.0")
    
    # 显示当前用户信息
    role_display_map = {"Sales": "销售员", "Prod": "生产/仓管", "Boss": "老板/管理", "Admin": "管理员"}
    current_role_display = role_display_map.get(st.session_state.role, st.session_state.role)
    
    st.success(f"👤 {st.session_state.operator_name}")
    st.caption(f"角色: {current_role_display}")
    
    if st.button("🚪 退出登录", use_container_width=True):
        st.session_state.current_user = None
        st.session_state.role = None
        st.session_state.is_admin = False
        st.rerun()
    
    st.divider()
    
    # 根据角色显示侧边栏功能
    user_perms = st.session_state.get('permissions', [])
    
    # 1. Management
    if "PLANNING" in user_perms:
        st.subheader("👑 管理功能")
        if st.button("👑 生产统筹/订单规划", use_container_width=True):
            go_page('boss_planning')
            st.rerun()
            
    # Admin User Management
    if st.session_state.role == "Admin":
        if st.button("👥 用户管理 (管理员)", use_container_width=True):
            go_page('user_management')
            st.rerun()
            
    # 2. Contract
    if "CONTRACT" in user_perms:
        if st.button("🏭 合同管理", use_container_width=True):
            go_page('production')
            st.rerun()

    st.divider()
    if st.button("📁 交易日志"): go_page('log_viewer')
    if "ARCHIVE" in user_perms:
        if st.button("📂 机台档案"): go_page('machine_archive')



# --- 🏠 首页 ---
if st.session_state.page == 'home':
    st.title("🏭 成品整机管理系统 V7.0")
    st.caption(f"当前用户: {st.session_state.operator_name} | 角色: {st.session_state.role}")
    st.write("")
    
    user_perms = st.session_state.get('permissions', [])
    
    # 1. 顶部：管理与统筹
    # 动态组装
    top_buttons = []
    if "PLANNING" in user_perms: top_buttons.append(FUNC_MAP["PLANNING"])
    if "CONTRACT" in user_perms: top_buttons.append(FUNC_MAP["CONTRACT"])
    
    if top_buttons:
        st.markdown("#### 👑 管理与统筹")
        c_adm = st.columns(len(top_buttons))
        for idx, btn in enumerate(top_buttons):
            with c_adm[idx]:
                st.markdown(f'<div class="{btn["class"]}">', unsafe_allow_html=True)
                if st.button(btn["label"], key=f"home_top_{btn['page']}", use_container_width=True): 
                    go_page(btn["page"]); st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
        st.divider()

    # 2. 核心业务按钮
    # 按照业务流顺序: INBOUND -> SALES_CREATE -> SALES_ALLOC -> SHIP_CONFIRM -> QUERY -> MACHINE_EDIT
    flow_order = ["INBOUND", "SALES_CREATE", "SALES_ALLOC", "SHIP_CONFIRM", "QUERY", "MACHINE_EDIT", "ARCHIVE"]
    
    core_buttons = []
    for code in flow_order:
        if code in user_perms and code not in ["PLANNING", "CONTRACT"]: # Avoid duplicates if any overlap
             core_buttons.append(FUNC_MAP[code])
        
    # 动态渲染按钮 (每行3个)
    if core_buttons:
        for i in range(0, len(core_buttons), 3):
            cols = st.columns(3, gap="medium")
            batch = core_buttons[i:i+3]
            for idx, btn in enumerate(batch):
                with cols[idx]:
                    st.markdown(f'<div class="big-btn {btn["class"]}">', unsafe_allow_html=True)
                    if st.button(btn["label"], key=f"home_btn_{btn['page']}"): 
                        go_page(btn["page"]); st.rerun()
                    st.markdown('</div>', unsafe_allow_html=True)

# --- 👥 用户管理 (Admin) ---
elif st.session_state.page == 'user_management':
    # Strict Admin Check - ONLY Admin role can access this, ignoring permissions
    if st.session_state.role != "Admin":
        st.error("🚫 权限不足！此页面仅限系统管理员访问。")
        st.button("⬅️ 返回首页", on_click=go_home)
        st.stop()

    c_back, c_title = st.columns([2, 8])
    with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
    with c_title: st.header("👥 用户注册审核与管理")
    
    users_df = get_all_users()
    
    # Metrics
    total = len(users_df)
    pending = len(users_df[users_df['status'] == 'pending'])
    active = len(users_df[users_df['status'] == 'active'])
    
    c1, c2, c3 = st.columns(3)
    c1.metric("总用户数", total)
    c2.metric("🟢 活跃用户", active)
    c3.metric("🟠 待审核", pending)
    
    st.divider()
    
    tab_audit, tab_all = st.tabs(["🟠 待审核申请", "📋 所有用户列表"])
    
    with tab_audit:
        pending_df = users_df[users_df['status'] == 'pending'].copy()
        if pending_df.empty:
            st.info("暂无待审核的注册申请")
        else:
            for idx, row in pending_df.iterrows():
                with st.container(border=True):
                    c_info, c_act = st.columns([3, 1])
                    with c_info:
                        role_map_audit = {"Sales": "销售员", "Prod": "生产/仓管", "Boss": "老板/管理", "Admin": "管理员"}
                        role_cn = role_map_audit.get(row['role'], row['role'])
                        st.markdown(f"**申请人:** {row['name']} (`{row['username']}`)")
                        st.caption(f"申请角色: **{role_cn}** | 申请时间: {row['register_time']}")
                    
                    with c_act:
                        c_a1, c_a2 = st.columns(2)
                        with c_a1:
                            if st.button("✅ 通过", key=f"app_{row['username']}"):
                                users_df.loc[users_df['username'] == row['username'], 'status'] = 'active'
                                users_df.loc[users_df['username'] == row['username'], 'auditor'] = st.session_state.operator_name
                                users_df.loc[users_df['username'] == row['username'], 'audit_time'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                save_all_users(users_df)
                                st.success(f"已批准 {row['name']}")
                                time.sleep(0.5); st.rerun()
                        with c_a2:
                            if st.button("❌ 拒绝", key=f"rej_{row['username']}"):
                                users_df.loc[users_df['username'] == row['username'], 'status'] = 'rejected'
                                users_df.loc[users_df['username'] == row['username'], 'auditor'] = st.session_state.operator_name
                                users_df.loc[users_df['username'] == row['username'], 'audit_time'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                save_all_users(users_df)
                                st.warning(f"已拒绝 {row['name']}")
                                time.sleep(0.5); st.rerun()

    with tab_all:
        # Show all users
        st.dataframe(
            users_df[['username', 'name', 'role', 'status', 'register_time', 'audit_time', 'auditor']],
            use_container_width=True,
            hide_index=True
        )

# --- 👑 生产统筹/订单规划 (老板核心功能) ---
elif st.session_state.page == 'boss_planning':
    check_access('PLANNING')

    if True:
        c_back, c_title = st.columns([1, 9])
        with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
        with c_title: st.header("👑 生产统筹 & 订单资源分配")

        # 恢复左右分栏布局
        col_list, col_detail = st.columns([1, 2], gap="large")

        fp_df = get_factory_plan()
        orders_df = get_orders()
        inventory_df = get_data()

        if 'boss_selected_id' not in st.session_state: st.session_state.boss_selected_id = None
        if 'boss_selected_type' not in st.session_state: st.session_state.boss_selected_type = 'contract' # contract | order

        # ==================== 左侧列表 (导航) ====================
        with col_list:
            with st.container(height=750, border=True):
                tab_pending, tab_planning, tab_orders = st.tabs(["📄 待审合同", "🎯 待规划", "📦 现有订单"])
                
                # --- 1. 待审合同 (Pending Review) ---
                with tab_pending:
                    # 搜索框
                    search_txt = st.text_input("🔍 搜索待审合同 (合同号/客户)", key="search_pending")
                    
                    # 按合同号聚合
                    pending_df = fp_df[fp_df['状态'] == '未下单'].copy()
                    known_contract_ids = pending_df['合同号'].dropna().astype(str).tolist()
                    orphan_df = get_unlinked_contract_folders(known_contract_ids)
                    
                    if search_txt:
                        s_term = search_txt.lower()
                        pending_df = pending_df[
                            pending_df['合同号'].str.lower().str.contains(s_term, na=False) |
                            pending_df['客户名'].str.lower().str.contains(s_term, na=False)
                        ]
                        if not orphan_df.empty:
                            orphan_df = orphan_df[
                                orphan_df['合同号'].astype(str).str.lower().str.contains(s_term, na=False) |
                                orphan_df['客户名'].astype(str).str.lower().str.contains(s_term, na=False)
                            ]
                    
                    if pending_df.empty and orphan_df.empty:
                        st.info("无待审合同")
                    else:
                        if not pending_df.empty:
                            pending_df['temp_date'] = pd.to_datetime(pending_df['要求交期'], errors='coerce')
                            pending_df = pending_df.sort_values('temp_date')
                            pending_df['month_key'] = pending_df['temp_date'].apply(lambda x: x.strftime('%Y-%m') if pd.notnull(x) else 'Unknown')
                            months = sorted([m for m in pending_df['month_key'].unique() if m != 'Unknown'], reverse=False)
                            if 'Unknown' in pending_df['month_key'].unique():
                                months.append('Unknown')

                            for m_key in months:
                                m_rows = pending_df[pending_df['month_key'] == m_key]
                                unique_contracts = m_rows['合同号'].unique()
                                count = len(unique_contracts)
                                is_expanded = (m_key == months[0])

                                with st.expander(f"📅 {m_key} ({count} 单)", expanded=is_expanded):
                                    for cid in unique_contracts:
                                        c_rows = m_rows[m_rows['合同号'] == cid]
                                        cust = c_rows.iloc[0]['客户名']
                                        model_counts = c_rows.groupby('机型')['排产数量'].apply(lambda x: sum(int(float(i)) for i in x))
                                        models_display = []
                                        for m, q in model_counts.items():
                                            models_display.append(f"{m} x{q}")
                                        models_str = "\n".join(models_display)

                                        label = f"🏢 {cust}\n{models_str}"
                                        btn_type = "primary" if (st.session_state.boss_selected_id == cid and st.session_state.boss_selected_type == 'contract') else "secondary"

                                        if st.button(label, key=f"btn_con_{cid}_{m_key}", type=btn_type, use_container_width=True):
                                            st.session_state.boss_selected_id = cid
                                            st.session_state.boss_selected_type = 'contract'
                                            st.rerun()

                        if not orphan_df.empty:
                            st.caption("📎 仅附件合同")
                            for idx, row in orphan_df.iterrows():
                                cid = str(row['合同号'])
                                cust = str(row['客户名'])
                                fcnt = int(row['文件数']) if str(row['文件数']).isdigit() else row['文件数']
                                ts = str(row['最近更新时间'])
                                label = f"📎 {cust}\n{cid} | {fcnt}个文件"
                                btn_type = "primary" if (st.session_state.boss_selected_id == cid and st.session_state.boss_selected_type == 'orphan_contract') else "secondary"
                                if st.button(label, key=f"btn_orphan_{cid}_{idx}", type=btn_type, use_container_width=True, help=f"最近更新: {ts}"):
                                    st.session_state.boss_selected_id = cid
                                    st.session_state.boss_selected_type = 'orphan_contract'
                                    st.rerun()

                # --- 2. 待规划 (Pending Planning) ---
                with tab_planning:
                    search_plan = st.text_input("🔍 搜索待规划 (合同/客户)", key="search_planning")
                    # 状态为 '待规划'
                    planning_df = fp_df[fp_df['状态'] == '待规划'].copy()

                    if search_plan:
                        s_term = search_plan.lower()
                        planning_df = planning_df[
                            planning_df['合同号'].str.lower().str.contains(s_term, na=False) |
                            planning_df['客户名'].str.lower().str.contains(s_term, na=False)
                        ]

                    if planning_df.empty:
                        st.info("无待规划项")
                    else:
                        # 1. 预处理数据和排序
                        planning_df['temp_date'] = pd.to_datetime(planning_df['要求交期'], errors='coerce')
                        planning_df = planning_df.sort_values('temp_date')
                        
                        # 2. 增加月份分组键
                        planning_df['month_key'] = planning_df['temp_date'].apply(lambda x: x.strftime('%Y-%m') if pd.notnull(x) else 'Unknown')
                        
                        # 3. 获取所有月份并排序
                        months = sorted([m for m in planning_df['month_key'].unique() if m != 'Unknown'], reverse=True)
                        if 'Unknown' in planning_df['month_key'].unique():
                            months.append('Unknown')
                        
                        # 4. 按月份渲染折叠面板
                        for m_key in months:
                            m_rows = planning_df[planning_df['month_key'] == m_key]
                            unique_plans = m_rows['合同号'].unique()
                            count = len(unique_plans)
                            
                            # 默认展开最近一个月
                            is_expanded = (m_key == months[0])
                            
                            with st.expander(f"📅 {m_key} ({count} 单)", expanded=is_expanded):
                                for cid in unique_plans:
                                    c_rows = m_rows[m_rows['合同号'] == cid]
                                    cust = c_rows.iloc[0]['客户名']
                                    models_summary = ", ".join(c_rows['机型'].unique())
                                    if len(models_summary) > 15: models_summary = models_summary[:15] + "..."
                                    
                                    label = f"🎯 {cid}\n{cust} | {models_summary}"
                                    btn_type = "primary" if (st.session_state.boss_selected_id == cid and st.session_state.boss_selected_type == 'planning') else "secondary"
                                    
                                    if st.button(label, key=f"btn_plan_{cid}_{m_key}", type=btn_type, use_container_width=True):
                                        st.session_state.boss_selected_id = cid
                                        st.session_state.boss_selected_type = 'planning'
                                        st.rerun()

                # --- 3. 现有订单 (Existing Orders / Planned) ---
                with tab_orders:
                    # 辅助函数：判断订单是否已完结 (Shipping >= Request)
                    # 计算所有占用 (Allocated + Shipped + Pending)
                    if not inventory_df.empty:
                        shipped_stats = inventory_df[inventory_df['状态'] == '已出库'].groupby('占用订单号').size().to_dict()
                        all_allocated_stats = inventory_df[inventory_df['占用订单号'] != ""].groupby('占用订单号').size().to_dict()
                    else: 
                        shipped_stats = {}
                        all_allocated_stats = {}

                    def is_order_completed(oid, req_qty_str):
                        if not oid: return False
                        s_qty = shipped_stats.get(oid, 0)
                        try: r_qty = int(float(req_qty_str))
                        except: r_qty = 999999
                        return s_qty >= r_qty and r_qty > 0
                    
                    # 辅助：判断是否已配齐 (Allocated >= Request)
                    def is_fully_allocated(oid, req_qty_str):
                        if not oid: return False
                        alloc_qty = all_allocated_stats.get(oid, 0)
                        try: r_qty = int(float(req_qty_str))
                        except: r_qty = 999999
                        return alloc_qty >= r_qty and r_qty > 0

                    # 1. 准备合同数据 (Contracts)
                    done_df = fp_df[fp_df['状态'].isin(['已规划', '已转订单', '已下单', '已配货'])].copy()
                    
                    # 过滤掉已完结的合同 (如果关联了订单且订单已完结)
                    completed_oids = set()
                    if not orders_df.empty and not done_df.empty:
                        linked_oids = done_df['订单号'].dropna().unique()
                        for oid in linked_oids:
                            if not oid: continue
                            ord_row = orders_df[orders_df['订单号'] == oid]
                            if not ord_row.empty:
                                req = ord_row.iloc[0]['需求数量']
                                if is_order_completed(oid, req):
                                    completed_oids.add(oid)
                    
                    if completed_oids:
                        done_df = done_df[~done_df['订单号'].isin(completed_oids)]

                    # 2. 准备手动订单数据 (Manual Orders - No Contract)
                    # 找出所有在 fp_df 中出现的订单号
                    all_contract_oids = set(fp_df['订单号'].dropna().unique())
                    
                    manual_orders = []
                    if not orders_df.empty:
                        for idx, row in orders_df.iterrows():
                            oid = row['订单号']
                            if not oid: continue
                            # 排除已关联合同的
                            if oid in all_contract_oids: continue
                            # 排除已完结的
                            if is_order_completed(oid, row['需求数量']): continue
                            
                            manual_orders.append(row)
                    
                    manual_df = pd.DataFrame(manual_orders)

                    search_q = st.text_input("🔍 搜索", placeholder="合同/订单/客户")
                    
                    # 搜索过滤
                    if search_q:
                        s = search_q.lower()
                        if not done_df.empty:
                            done_df = done_df[
                                done_df['合同号'].str.lower().str.contains(s, na=False) |
                                done_df['客户名'].str.lower().str.contains(s, na=False) |
                                done_df['订单号'].str.lower().str.contains(s, na=False)
                            ]
                        if not manual_df.empty:
                            manual_df = manual_df[
                                manual_df['订单号'].str.lower().str.contains(s, na=False) |
                                manual_df['客户名'].str.lower().str.contains(s, na=False) |
                                manual_df['代理商'].str.lower().str.contains(s, na=False)
                            ]

                    # 显示列表
                    has_data = False
                    
                    # A. 合同列表 (按月份折叠)
                    if not done_df.empty:
                        has_data = True
                        st.caption("📑 合同订单 (按月分组)")
                        
                        # 1. 增加月份辅助列 (基于 '要求交期')
                        done_df['month_key'] = pd.to_datetime(done_df['要求交期'], errors='coerce').dt.strftime('%Y-%m')
                        done_df['month_key'] = done_df['month_key'].fillna('Unknown')
                        
                        # 2. 获取所有月份并降序排列
                        all_months = sorted(done_df['month_key'].unique(), reverse=True)
                        
                        for m_key in all_months:
                            # 3. 统计该月合同数
                            m_rows = done_df[done_df['month_key'] == m_key]
                            m_count = m_rows['合同号'].nunique()
                            
                            # 4. 默认展开最近一个月，其他折叠
                            is_expanded = (m_key == all_months[0])
                            
                            with st.expander(f"📅 {m_key} ({m_count} 单)", expanded=is_expanded):
                                unique_done = m_rows['合同号'].unique()[::-1]
                                for cid in unique_done:
                                    c_rows = m_rows[m_rows['合同号'] == cid]
                                    cust = c_rows.iloc[0]['客户名']
                                    oid = c_rows.iloc[0].get('订单号', '')
                                    status = c_rows.iloc[0]['状态']
                                    
                                label = f"📦 {cid} ({status})\n{cust}" + (f" | {oid}" if oid else "")
                                btn_type = "primary" if (st.session_state.boss_selected_id == cid and st.session_state.boss_selected_type == 'done') else "secondary"
                                    
                                    # 👇 在这里加上 _{m_key}
                                if st.button(label, key=f"btn_done_{cid}_{m_key}", type=btn_type, use_container_width=True):
                                        st.session_state.boss_selected_id = cid
                                        st.session_state.boss_selected_type = 'done'
                                        st.rerun()
                    
                    # B. 手动订单列表
                    if not manual_df.empty:
                        has_data = True
                        st.caption("📝 独立订单 (无合同)")
                        
                        # 分组：进行中 vs 已配齐
                        m_ongoing = []
                        m_allocated = []
                        
                        for idx, row in manual_df.iloc[::-1].iterrows():
                            if is_fully_allocated(row['订单号'], row['需求数量']):
                                m_allocated.append(row)
                            else:
                                m_ongoing.append(row)
                        
                        # 渲染函数
                        def render_manual_btn(row):
                            oid = row['订单号']
                            cust = row['客户名']
                            s_qty = shipped_stats.get(oid, 0)
                            req_qty = row['需求数量']
                            label = f"📝 {oid}\n{cust} | 进度: {s_qty}/{req_qty}"
                            btn_type = "primary" if (st.session_state.boss_selected_id == oid and st.session_state.boss_selected_type == 'manual_order') else "secondary"
                            if st.button(label, key=f"btn_manual_{oid}", type=btn_type, use_container_width=True):
                                st.session_state.boss_selected_id = oid
                                st.session_state.boss_selected_type = 'manual_order'
                                st.rerun()

                        # 1. 显示进行中
                        for row in m_ongoing: render_manual_btn(row)
                        
                        # 2. 显示已配齐 (折叠)
                        if m_allocated:
                            with st.expander(f"✅ 已配齐待发 ({len(m_allocated)})", expanded=False):
                                for row in m_allocated: render_manual_btn(row)

                    if not has_data:
                        st.info("无相关未完结订单")

        # ==================== 右侧详情 (Review / Plan / Edit) ====================
        with col_detail:
            with st.container(height=750, border=True):
                sel_id = st.session_state.boss_selected_id
                sel_type = st.session_state.boss_selected_type
                
                if not sel_id:
                     # 用户要求：右边配置区域也是默认空白
                     # 原来是: st.info("👈 请从左侧选择一个项目")
                     st.info("👈 请从左侧选择一个项目以查看详情")
                
                # --- 场景1: 待审合同 (Review) ---
                elif sel_type == 'contract':
                    target_rows = fp_df[(fp_df['合同号'] == sel_id) & (fp_df['状态'] == '未下单')]
                    if target_rows.empty:
                        st.warning("该合同状态已变更，请刷新")
                    else:
                        first_row = target_rows.iloc[0]
                        
                        # --- 编辑模式开关 ---
                        c_h1, c_h2 = st.columns([8, 2])
                        with c_h1: st.markdown(f"### 📄 合同详情: {sel_id}")
                        with c_h2: 
                            is_edit_mode = st.toggle("✏️ 编辑模式", key=f"edit_mode_{sel_id}")
                        
                        if is_edit_mode:
                            st.info("正在编辑合同内容...")
                            with st.form(f"edit_contract_{sel_id}"):
                                c_e1, c_e2 = st.columns(2)
                                with c_e1:
                                    new_cust = st.text_input("客户名", value=first_row['客户名'])
                                    new_deadline = st.date_input("要求交期", value=pd.to_datetime(first_row['要求交期']).date() if pd.notna(first_row['要求交期']) else datetime.now().date())
                                with c_e2:
                                    new_agent = st.text_input("代理商", value=first_row['代理商'])
                                    # 备注通常存在第一行或者每行都有，这里取第一行作为总备注的近似，或者允许批量修改
                                    # 实际上 factory_plan 里的备注是行备注。但通常合同有个总备注？
                                    # 我们的数据结构里没有独立的合同总备注，通常存在每行的备注里，或者只是UI上的概念
                                    # 这里我们提供一个 "批量修改备注" 的功能，或者只展示行编辑
                                    pass

                                st.markdown("#### 🛠️ 机型明细 (可增删改)")
                                # 准备编辑用的数据
                                # 必须包含: 机型, 排产数量, 备注
                                edit_data = target_rows[['机型', '排产数量', '备注']].copy()
                                # 转换数量为数字
                                edit_data['排产数量'] = edit_data['排产数量'].astype(float).astype(int)
                                
                                # 获取全量机型选项
                                df_inv_tmp = get_data()
                                all_models = set(CUSTOM_MODEL_ORDER)
                                if not df_inv_tmp.empty: all_models.update(df_inv_tmp['机型'].unique())
                                model_options = sorted(list(all_models), key=get_model_rank)

                                edited_items = st.data_editor(
                                    edit_data,
                                    num_rows="dynamic",
                                    use_container_width=True,
                                    column_config={
                                        "机型": st.column_config.SelectboxColumn("机型", options=model_options, required=True),
                                        "排产数量": st.column_config.NumberColumn("数量", min_value=1, required=True),
                                        "备注": st.column_config.TextColumn("备注")
                                    }
                                )
                                
                                if st.form_submit_button("💾 保存修改"):
                                    if edited_items.empty:
                                        st.error("至少需要一行机型数据")
                                    else:
                                        # 1. 删除旧数据
                                        fp_df = fp_df[fp_df['合同号'] != sel_id]
                                        
                                        # 2. 构建新数据
                                        new_rows = []
                                        for _, r_item in edited_items.iterrows():
                                            new_row = {
                                                "合同号": sel_id,
                                                "机型": r_item['机型'],
                                                "排产数量": int(r_item['排产数量']),
                                                "要求交期": str(new_deadline),
                                                "状态": "未下单",
                                                "备注": r_item.get('备注', ''),
                                                "客户名": new_cust,
                                                "代理商": new_agent,
                                                "指定批次/来源": "", # 重置规划
                                                "订单号": ""
                                            }
                                            new_rows.append(new_row)
                                        
                                        fp_df = pd.concat([fp_df, pd.DataFrame(new_rows)], ignore_index=True)
                                        save_factory_plan(fp_df)
                                        st.success("合同修改已保存！")
                                        # 关闭编辑模式
                                        # st.session_state[f"edit_mode_{sel_id}"] = False # toggle 无法直接通过代码关闭，只能 rerun
                                        time.sleep(1); st.rerun()

                        else:
                            # --- 阅读模式 (原显示逻辑) ---
                            st.markdown(f"""
                            <div class="boss-plan-card">
                                <p><b>客户:</b> {first_row['客户名']} | <b>代理:</b> {first_row['代理商']}</p>
                                <p><b>交期:</b> {first_row['要求交期']}</p>
                                <p><b>备注:</b> {first_row['备注']}</p>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            st.write("**包含机型:**")
                            st.dataframe(target_rows[['机型', '排产数量', '要求交期', '备注']], use_container_width=True, hide_index=True)
                        
                        # --- V7.5 File Preview & Download (Enhanced UI) ---
                        render_file_manager(sel_id, first_row['客户名'], default_expanded=True)

                        st.divider()
                        c_act1, c_act2 = st.columns(2)
                        with c_act1:
                            if st.button("🚀 前往规划", type="primary", use_container_width=True):
                                # 将该合同下所有条目状态改为 '待规划'
                                fp_df.loc[fp_df['合同号'] == sel_id, '状态'] = '待规划'
                                save_factory_plan(fp_df)
                                st.session_state.boss_selected_type = 'planning'
                                st.success("已批准！进入规划阶段。")
                                st.rerun()
                        with c_act2:
                            if st.button("❌ 驳回/取消", use_container_width=True):
                                fp_df.loc[fp_df['合同号'] == sel_id, '状态'] = '已取消'
                                save_factory_plan(fp_df)
                                st.warning("合同已取消")
                                st.session_state.boss_selected_id = None
                                st.rerun()

                elif sel_type == 'orphan_contract':
                    orphan_id = str(sel_id)
                    orphan_customer = orphan_id.split("_")[0] if "_" in orphan_id else orphan_id
                    st.markdown(f"### 📎 附件详情: {orphan_id}")
                    st.caption("该条目来自 data 目录附件，尚未在合同计划中建档。")
                    render_file_manager(orphan_id, orphan_customer, default_expanded=True, key_suffix="_orphan")

                # --- 场景2: 待规划 & 已规划 (Planning / Editing) ---
                elif sel_type in ['planning', 'done']:
                    # 获取该合同所有行
                    target_rows = fp_df[fp_df['合同号'] == sel_id]
                    if target_rows.empty:
                        st.error("数据未找到")
                    else:
                        first_row = target_rows.iloc[0]
                        status_now = first_row['状态']
                        oid_now = first_row.get('订单号', '')
                        
                        st.markdown(f"""
                        <div class="boss-plan-card">
                            <h3>🎯 规划详情: {sel_id}</h3>
                            <p><b>状态:</b> {status_now} {f'| <b>关联订单:</b> {oid_now}' if oid_now else ''}</p>
                            <p><b>客户:</b> {first_row['客户名']} | <b>交期:</b> {first_row['要求交期']}</p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # --- V7.7 Add File Preview (Collapsed) ---
                        render_file_manager(sel_id, first_row['客户名'], default_expanded=False, key_suffix="_plan")

                        changes_map = {} # {idx: new_plan_str}
                        
                        for idx, row in target_rows.iterrows():
                            model = row['机型']
                            qty_needed = int(float(row['排产数量']))
                            remark = row.get('备注', '')
                            
                            # --- 解析机型是否为加高 ---
                            is_high_req = "(加高)" in model
                            real_model = model.replace("(加高)", "")
                            req_label = f"🏗️ {real_model} [加高定制]" if is_high_req else f"⚙️ {real_model} [标准]"

                            st.markdown(f"#### {req_label} (需 {qty_needed} 台)")
                            if remark:
                                st.info(f"📝 **备注:** {remark}")
                            
                            # 解析旧配置
                            saved_plan_str = str(row.get('指定批次/来源', ''))
                            prev_alloc = {}
                            try:
                                 if ":" in saved_plan_str and not saved_plan_str.strip().startswith("{"):
                                     prev_alloc = ast.literal_eval(saved_plan_str.split(":", 1)[1])
                                 else:
                                     prev_alloc = ast.literal_eval(saved_plan_str)
                            except: pass
                            
                            # --- 新增：显示实际配货进度 (Actual Allocation) ---
                            actual_alloc_info = ""
                            has_actual = False
                            if oid_now:
                                # 查找 inventory 中被此订单占用的
                                act_mask = (inventory_df['占用订单号'] == oid_now) & (inventory_df['机型'] == real_model)
                                if is_high_req:
                                    act_mask = act_mask & (inventory_df['机台备注/配置'].str.contains("加高", na=False))
                                else:
                                    act_mask = act_mask & (~inventory_df['机台备注/配置'].str.contains("加高", na=False))
                                
                                act_df = inventory_df[act_mask]
                                
                                if not act_df.empty:
                                    has_actual = True
                                    act_counts = act_df.groupby('批次号').size().to_dict()
                                    info_parts = []
                                    for b, c in act_counts.items():
                                        info_parts.append(f"{b}: {c}台")
                                    actual_alloc_info = f"✅ **实际已配:** {', '.join(info_parts)} (共 {len(act_df)} 台)"
                                else:
                                    actual_alloc_info = "ℹ️ 实际未配货"
                            
                            if has_actual:
                                # 恢复直观展示
                                st.write(actual_alloc_info)
                                # 同步按钮
                                if st.button("📥 载入实际配货到规划", key=f"sync_{idx}_{model}", help="将当前实际配货数量填入下方规划框"):
                                    # 构造新的配货字典
                                    new_alloc = {}
                                    for b, c in act_counts.items():
                                        if b == "无批次" or not b: new_alloc['现货(Spot)'] = c
                                        else: new_alloc[b] = c
                                    
                                    # 1. 更新当前行的 '指定批次/来源'
                                    new_plan_str = f"{model}:{str(new_alloc)}"
                                    fp_df.loc[idx, '指定批次/来源'] = new_plan_str
                                    save_factory_plan(fp_df)
                                    
                                    # 2. 如果关联了订单，同步更新 Sales Order
                                    if oid_now:
                                        # 重新获取整个合同的 plan
                                        # 注意：这里只更新了当前 model 的 plan，其他 model 保持原样
                                        # 为了确保完整性，我们从 fp_df 重新构建整个合同的 plan string
                                        # 但 fp_df 在内存中可能还没包含其他 model 的最新改动（如果用户没点保存）
                                        # 这里是一个单点操作，我们假设只修改这一条
                                        
                                        # 获取该合同所有相关行（最新状态）
                                        related_rows = fp_df[fp_df['合同号'] == sel_id]
                                        all_plans_sync = []
                                        for _, r_sync in related_rows.iterrows():
                                            all_plans_sync.append(str(r_sync.get('指定批次/来源', '')))
                                        
                                        combined_plan_sync = "; ".join(all_plans_sync)
                                        
                                        if oid_now in orders_df['订单号'].values:
                                            orders_df.loc[orders_df['订单号'] == oid_now, '指定批次/来源'] = combined_plan_sync
                                            save_orders(orders_df)
                                    
                                    st.success("已同步实际配货数据！"); time.sleep(0.5); st.rerun()
                            else:
                                 st.caption(actual_alloc_info)

                            # 库存查询
                            # 1. Spot (Left): 仅包含 '库存中'
                            mask_spot = (
                                (inventory_df['机型'] == real_model) & 
                                (inventory_df['状态'] == '库存中') & 
                                (inventory_df['占用订单号'] == "")
                            )
                            if is_high_req:
                                mask_spot = mask_spot & (inventory_df['机台备注/配置'].str.contains("加高", na=False))
                            else:
                                mask_spot = mask_spot & (~inventory_df['机台备注/配置'].str.contains("加高", na=False))
                            
                            spot_df = inventory_df[mask_spot]
                            spot_count = len(spot_df)

                            # 2. Batch (Right): 仅包含 '待入库'
                            mask_batch = (
                                (inventory_df['机型'] == real_model) & 
                                (inventory_df['状态'] == '待入库') & 
                                (inventory_df['占用订单号'] == "")
                            )
                            if is_high_req:
                                mask_batch = mask_batch & (inventory_df['机台备注/配置'].str.contains("加高", na=False))
                            else:
                                mask_batch = mask_batch & (~inventory_df['机台备注/配置'].str.contains("加高", na=False))
                                
                            batch_df = inventory_df[mask_batch]
                            batch_stats = batch_df.groupby('批次号').size().to_dict()
                            
                            # --- 获取批次预计入库时间 ---
                            batch_dates = {}
                            if '预计入库时间' in batch_df.columns:
                                temp_dates = batch_df[['批次号', '预计入库时间']].drop_duplicates()
                                for _, r_d in temp_dates.iterrows():
                                    b_name = r_d['批次号']
                                    d_val = str(r_d['预计入库时间']).strip()
                                    if b_name not in batch_dates and d_val:
                                        batch_dates[b_name] = d_val
                            
                            c_r1, c_r2 = st.columns(2)
                            with c_r1:
                                val_spot = prev_alloc.get('现货(Spot)', 0)
                                alloc_spot = st.number_input(f"现货 (余{spot_count})", min_value=0, value=int(val_spot), key=f"plan_spot_{idx}")
                            
                            current_batch_alloc = {}
                            with c_r2:
                                all_batches = set(batch_stats.keys()) | set(k for k in prev_alloc.keys() if k != '现货(Spot)')
                                if not all_batches: st.caption("无批次库存")
                                else:
                                    for b in sorted(list(all_batches)):
                                        can_use = batch_stats.get(b, 0)
                                        
                                        # 构建显示标签
                                        label_text = f"批次 {b} (余{can_use})"
                                        arr_time = batch_dates.get(b, "")
                                        if arr_time:
                                            label_text += f" [📅 {arr_time}]"
                                        
                                        prev_val = prev_alloc.get(b, 0)
                                        val_b = st.number_input(label_text, min_value=0, value=int(prev_val), key=f"plan_b_{idx}_{b}")
                                        if val_b > 0: current_batch_alloc[b] = val_b
                            
                            total_alloc = alloc_spot + sum(current_batch_alloc.values())
                            st.progress(min(total_alloc / qty_needed, 1.0) if qty_needed > 0 else 0)
                            st.caption(f"已规划: {total_alloc} / {qty_needed}")
                            
                            # 构建保存串
                            this_plan = {}
                            if alloc_spot > 0: this_plan['现货(Spot)'] = alloc_spot
                            this_plan.update(current_batch_alloc)
                            
                            final_str = f"{model}:{str(this_plan)}"
                            changes_map[idx] = final_str
                            st.divider()

                        if st.button("💾 保存规划 (Save Plan)", type="primary"):
                            for idx, plan_str in changes_map.items():
                                fp_df.loc[idx, '指定批次/来源'] = plan_str
                                
                                # --- Save to CSV Record ---
                                if oid_now:
                                    try:
                                        model_name = fp_df.loc[idx, '机型']
                                        val_to_save = plan_str
                                        if ":" in plan_str:
                                            _, val_to_save = plan_str.split(":", 1)
                                        save_planning_record(oid_now, model_name, val_to_save)
                                    except Exception as e:
                                        print(f"Error saving to CSV: {e}")

                                if fp_df.loc[idx, '状态'] == '待规划':
                                    fp_df.loc[idx, '状态'] = '已规划'
                            
                            save_factory_plan(fp_df)
                            
                            if oid_now:
                                all_plans = []
                                for idx, plan_str in changes_map.items():
                                    all_plans.append(plan_str)
                                combined_plan_str = "; ".join(all_plans)
                                if oid_now in orders_df['订单号'].values:
                                    orders_df.loc[orders_df['订单号'] == oid_now, '指定批次/来源'] = combined_plan_str
                                    save_orders(orders_df)
                                    st.success(f"已同步更新销售订单 {oid_now}！")
                                else:
                                    st.warning(f"关联订单 {oid_now} 在销售表中未找到，仅更新了规划表。")
                            else:
                                st.success("规划已保存！(等待销售下单引用)")
                            
                            time.sleep(1); st.rerun()
                        
                        if status_now == '已规划' and not oid_now:
                            if st.button("🚀 直通配货 (自动生成销售订单)", help="跳过销售确认，直接生成订单并进入配货"):
                                model_data = {}
                                all_plans = []
                                note_combined = ""
                                for idx, row in target_rows.iterrows():
                                    m = row['机型']; q = int(float(row['排产数量']))
                                    model_data[m] = q
                                    if row.get('备注'): note_combined += f" {m}:{row['备注']}"
                                    all_plans.append(str(row.get('指定批次/来源', '')))
                                combined_plan_str = "; ".join(all_plans)
                                new_oid = create_sales_order(
                                    customer=first_row['客户名'], agent=first_row['代理商'], model_data=model_data,
                                    note=note_combined, pack_option="未指定", delivery_time=first_row['要求交期'],
                                    source_batch=combined_plan_str
                                )
                                fp_df.loc[fp_df['合同号'] == sel_id, '订单号'] = new_oid
                                fp_df.loc[fp_df['合同号'] == sel_id, '状态'] = '已转订单'
                                save_factory_plan(fp_df)
                                st.success(f"已自动生成订单 {new_oid}！"); time.sleep(1); st.rerun()

                # --- 场景3: 独立手动订单 (Manual Orders) ---
                elif sel_type == 'manual_order':
                    target_order = orders_df[orders_df['订单号'] == sel_id]
                    if target_order.empty:
                        st.error("订单未找到")
                    else:
                        row = target_order.iloc[0]
                        cust = row['客户名']; agent = row['代理商']
                        
                        st.markdown(f"""
                        <div class="boss-plan-card">
                            <h3>📝 独立订单规划: {sel_id}</h3>
                            <p><b>客户:</b> {cust} | <b>代理:</b> {agent}</p>
                            <p><b>发货时间:</b> {row['发货时间']} | <b>备注:</b> {row['备注']}</p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # 解析需求机型
                        reqs = parse_requirements(row['需求机型'], row['需求数量'])
                        
                        # 解析已有的 source plan
                        existing_plan_str = str(row.get('指定批次/来源', ''))
                        existing_plan_map = {} # {Model: {Batch: Qty}}
                        if existing_plan_str:
                            try:
                                parts = existing_plan_str.split(";")
                                for p in parts:
                                    if ":" in p:
                                        m, content = p.split(":", 1)
                                        m = m.strip()
                                        try: existing_plan_map[m] = ast.literal_eval(content.strip())
                                        except: pass
                            except: pass

                        new_plans = []
                        
                        for model_key, qty in reqs.items():
                            # --- V7.1 核心解析逻辑 ---
                            is_high_req = "(加高)" in model_key
                            real_model = model_key.replace("(加高)", "")
                            req_label = f"🏗️ {real_model} [加高定制]" if is_high_req else f"⚙️ {real_model} [标准]"
                            
                            st.markdown(f"#### {req_label} (需 {qty} 台)")
                            
                            prev_alloc = existing_plan_map.get(model_key, {})
                            
                            # --- 新增：显示实际配货进度 (Actual Allocation) for Manual Orders ---
                            actual_alloc_info = ""
                            has_actual = False
                            
                            # 查找 inventory 中被此订单占用的
                            # V7.1: 使用 real_model 并结合备注过滤
                            act_mask = (inventory_df['占用订单号'] == sel_id) & (inventory_df['机型'] == real_model)
                            if is_high_req:
                                act_mask = act_mask & (inventory_df['机台备注/配置'].str.contains("加高", na=False))
                            else:
                                act_mask = act_mask & (~inventory_df['机台备注/配置'].str.contains("加高", na=False))
                                
                            act_df = inventory_df[act_mask]
                            
                            if not act_df.empty:
                                has_actual = True
                                act_counts = act_df.groupby('批次号').size().to_dict()
                                info_parts = []
                                for b, c in act_counts.items():
                                    info_parts.append(f"{b}: {c}台")
                                actual_alloc_info = f"✅ **实际已配:** {', '.join(info_parts)} (共 {len(act_df)} 台)"
                            else:
                                actual_alloc_info = "ℹ️ 实际未配货"
                            
                            if has_actual:
                                # 恢复直观展示
                                st.write(actual_alloc_info)
                                # 同步按钮
                                if st.button("📥 载入实际配货到规划", key=f"sync_m_{sel_id}_{model_key}", help="将当前实际配货数量填入下方规划框"):
                                    # 构造新的配货字典
                                    new_alloc = {}
                                    for b, c in act_counts.items():
                                        if b == "无批次" or not b: new_alloc['现货(Spot)'] = c
                                        else: new_alloc[b] = c
                                    
                                    # 更新 Manual Order 的指定批次/来源
                                    # 1. 更新内存 map
                                    existing_plan_map[model_key] = new_alloc
                                    
                                    # 2. 重新构建整个订单的 plan string
                                    # Manual Order 的格式是 "ModelA:{...}; ModelB:{...}"
                                    all_plans_sync = []
                                    # 注意：reqs 是 {model: qty}，我们遍历 reqs 来重组字符串，确保覆盖所有机型
                                    for m_key in reqs.keys():
                                        alloc_data = existing_plan_map.get(m_key, {})
                                        if alloc_data:
                                            all_plans_sync.append(f"{m_key}:{str(alloc_data)}")
                                    
                                    combined_plan_sync = "; ".join(all_plans_sync)
                                    
                                    orders_df.loc[orders_df['订单号'] == sel_id, '指定批次/来源'] = combined_plan_sync
                                    save_orders(orders_df)
                                    st.success("已同步实际配货数据！"); time.sleep(0.5); st.rerun()
                            else:
                                 st.caption(actual_alloc_info)

                            # 库存查询
                            # 1. Spot (Left): 仅包含 '库存中'
                            mask_spot = (
                                (inventory_df['机型'] == real_model) & 
                                (inventory_df['状态'] == '库存中') & 
                                (inventory_df['占用订单号'] == "")
                            )
                            if is_high_req:
                                mask_spot = mask_spot & (inventory_df['机台备注/配置'].str.contains("加高", na=False))
                            else:
                                mask_spot = mask_spot & (~inventory_df['机台备注/配置'].str.contains("加高", na=False))
                            
                            spot_df = inventory_df[mask_spot]
                            spot_count = len(spot_df)

                            # 2. Batch (Right): 仅包含 '待入库'
                            mask_batch = (
                                (inventory_df['机型'] == real_model) & 
                                (inventory_df['状态'] == '待入库') & 
                                (inventory_df['占用订单号'] == "")
                            )
                            if is_high_req:
                                mask_batch = mask_batch & (inventory_df['机台备注/配置'].str.contains("加高", na=False))
                            else:
                                mask_batch = mask_batch & (~inventory_df['机台备注/配置'].str.contains("加高", na=False))
                                
                            batch_df = inventory_df[mask_batch]
                            batch_stats = batch_df.groupby('批次号').size().to_dict()
                            
                            # --- 获取批次预计入库时间 ---
                            batch_dates = {}
                            if '预计入库时间' in batch_df.columns:
                                temp_dates = batch_df[['批次号', '预计入库时间']].drop_duplicates()
                                for _, r_d in temp_dates.iterrows():
                                    b_name = r_d['批次号']
                                    d_val = str(r_d['预计入库时间']).strip()
                                    if b_name not in batch_dates and d_val:
                                        batch_dates[b_name] = d_val
                            
                            c_r1, c_r2 = st.columns(2)
                            with c_r1:
                                val_spot = prev_alloc.get('现货(Spot)', 0)
                                alloc_spot = st.number_input(f"现货 (余{spot_count})", min_value=0, value=int(val_spot), key=f"m_spot_{model_key}")
                            
                            current_batch_alloc = {}
                            with c_r2:
                                all_batches = set(batch_stats.keys()) | set(k for k in prev_alloc.keys() if k != '现货(Spot)')
                                if not all_batches: st.caption("无批次库存")
                                else:
                                    for b in sorted(list(all_batches)):
                                        can_use = batch_stats.get(b, 0)
                                        
                                        # 构建显示标签
                                        label_text = f"批次 {b} (余{can_use})"
                                        arr_time = batch_dates.get(b, "")
                                        if arr_time:
                                            label_text += f" [📅 {arr_time}]"
                                        
                                        prev_val = prev_alloc.get(b, 0)
                                        val_b = st.number_input(label_text, min_value=0, value=int(prev_val), key=f"m_b_{model_key}_{b}")
                                        if val_b > 0: current_batch_alloc[b] = val_b
                            
                            total_alloc = alloc_spot + sum(current_batch_alloc.values())
                            st.progress(min(total_alloc / qty, 1.0) if qty > 0 else 0)
                            st.caption(f"已规划: {total_alloc} / {qty}")
                            
                            this_plan = {}
                            if alloc_spot > 0: this_plan['现货(Spot)'] = alloc_spot
                            this_plan.update(current_batch_alloc)
                            new_plans.append(f"{model_key}:{str(this_plan)}")
                            st.divider()
                            
                        if st.button("💾 保存订单规划", type="primary"):
                            final_plan_str = "; ".join(new_plans)
                            orders_df.loc[orders_df['订单号'] == sel_id, '指定批次/来源'] = final_plan_str
                            save_orders(orders_df)

                            # --- Save to CSV ---
                            try:
                                for p in new_plans:
                                    if ":" in p:
                                        m, c = p.split(":", 1)
                                        save_planning_record(sel_id, m, c)
                            except Exception as e:
                                print(f"Error saving manual plan to CSV: {e}")

                            st.success("订单规划已保存！"); time.sleep(1); st.rerun()

# --- 🏭 合同管理 (Contract Management) ---
elif st.session_state.page == 'production':
    check_access('CONTRACT')
    if True:
        c_back, c_title = st.columns([2, 8])
        with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
        with c_title: st.header("🏭 合同管理")
        st.info("💡 在此录入未来合同。老板审批后，将流转至销售下单环节。")
        
        with st.expander("➕ 新增合同 (批量录入)", expanded=False):
            if 'contract_models' not in st.session_state: st.session_state.contract_models = []
            
            df_inv = get_data()
            
            # 使用全量机型 (0库存也要能选)
            all_known_models = set(CUSTOM_MODEL_ORDER)
            if not df_inv.empty:
                all_known_models.update(df_inv['机型'].unique())
            available_models_prod = sorted(list(all_known_models), key=get_model_rank)

            # 1. 基础信息
            c1, c2 = st.columns(2)
            with c1: 
                # f_contract = st.text_input("合同号 (Contract No)")
                st.markdown("##### 合同号将自动生成")
                st.caption("格式: HT + 日期 + 随机4位数")
            with c2: 
                # Auto-fill deadline
                def_date = datetime.now().date()
                if 'auto_deadline' in st.session_state:
                    try: def_date = pd.to_datetime(st.session_state['auto_deadline']).date()
                    except: pass
                f_deadline = st.date_input("要求交期", value=def_date)
            
            c3, c4 = st.columns(2)
            with c3: 
                # Auto-fill customer
                val_cust = st.session_state.get('auto_customer', "")
                f_customer = st.text_input("客户名 (Customer)", value=val_cust)
            with c4: 
                val_agent = st.session_state.get('auto_agent', "")
                f_agent = st.text_input("代理商 (Agent)", value=val_agent)
            
            # --- V7.5 Contract File Upload ---
            st.markdown("##### 📎 合同附件 (可选)")
            f_file = st.file_uploader("上传合同文件 (PDF/Word/JPG, Max 50MB)", type=['pdf', 'doc', 'docx', 'jpg', 'jpeg'])
            
            # --- V7.6 Intelligent Identification ---
            if f_file:
                if st.button("✨ 智能识别 (AI OCR)", type="primary"):
                    with st.spinner("正在分析合同内容..."):
                        processor = OCRProcessor()
                        res_data, full_text = processor.process_file(f_file)
                        
                        if res_data:
                            # 1. 基础信息自动填充
                            if res_data.get('customer'):
                                st.session_state['auto_customer'] = res_data['customer']
                            if res_data.get('agent'):
                                st.session_state['auto_agent'] = res_data['agent']
                            if res_data.get('global_note'):
                                st.session_state['auto_global_note'] = res_data['global_note']
                                
                            # 2. 日期解析
                            if res_data.get('deadline'):
                                try:
                                    # 尝试解析多种格式
                                    raw_date = res_data['deadline']
                                    d_str = raw_date.replace("年", "-").replace("月", "-").replace("日", "").replace("/", "-").strip()
                                    pd.to_datetime(d_str) # check validity
                                    st.session_state['auto_deadline'] = d_str
                                except: pass
                                
                            # 3. 机型列表解析 (适配新 JSON 结构)
                            new_entry_list = []
                            items = res_data.get('items', [])
                            
                            # 兼容旧格式 (如果 LLM 返回了旧结构)
                            if not items and res_data.get('机型及数量'):
                                # Fallback logic for old structure string "A:1 | B:2"
                                raw_str = res_data['机型及数量']
                                if raw_str != '未识别':
                                    parts = raw_str.split("|")
                                    for p in parts:
                                        m, q = p, "1"
                                        if ":" in p: m, q = p.split(":", 1)
                                        new_entry_list.append({
                                            "model": m.strip(), "qty": int(q) if q.isdigit() else 1, 
                                            "is_high": False, "note": "AI识别(旧格式)"
                                        })

                            # 处理新结构列表
                            elif isinstance(items, list):
                                for item in items:
                                    m_name = item.get('model', '')
                                    q_val = item.get('qty', 1)
                                    is_h = item.get('is_high', False)
                                    note_val = item.get('note', '')
                                    
                                    if not m_name: continue
                                    
                                    new_entry_list.append({
                                        "model": m_name, 
                                        "qty": int(q_val) if str(q_val).isdigit() else 1,
                                        "is_high": bool(is_h),
                                        "note": note_val
                                    })

                            # 4. 模糊匹配并填充表格
                            final_table_data = []
                            for entry in new_entry_list:
                                m = entry['model']
                                best_match = None
                                
                                # Fuzzy Match Logic
                                if available_models_prod:
                                    # 优先完全匹配
                                    if m in available_models_prod:
                                        best_match = m
                                    else:
                                        # 包含匹配
                                        for known in available_models_prod:
                                            if known in m or m in known:
                                                best_match = known
                                                break
                                
                                final_table_data.append({
                                    "机型": best_match if best_match else (m if m else available_models_prod[0]),
                                    "数量": entry['qty'],
                                    "加高?": entry['is_high'],
                                    "单行备注": entry['note'] if entry['note'] else "AI识别"
                                })
                            
                            if final_table_data:
                                st.session_state.contract_entry_df = pd.DataFrame(final_table_data)
                                st.success(f"已自动提取 {len(final_table_data)} 条机型数据！")
                            
                            st.success("识别完成！请检查下方填入的数据。")
                            st.expander("查看识别原文").write(full_text)
                        else:
                            st.error("识别失败或未提取到有效信息")
            
            st.divider()
            
            # 2. 机型选择与数据录入 (使用 data_editor 以支持多选重复机型)
            st.caption("请在下方表格中添加机型，支持同一机型添加多行（例如一行标准、一行加高）。")
            
            if 'contract_entry_df' not in st.session_state:
                st.session_state.contract_entry_df = pd.DataFrame(
                    [{"机型": available_models_prod[0] if available_models_prod else "", "数量": 1, "加高?": False, "备注": ""}]
                )

            edited_df = st.data_editor(
                st.session_state.contract_entry_df,
                num_rows="dynamic",
                use_container_width=True,
                column_config={
                    "机型": st.column_config.SelectboxColumn("机型", options=available_models_prod, required=True),
                    "数量": st.column_config.NumberColumn("数量", min_value=1, default=1, required=True),
                    "加高?": st.column_config.CheckboxColumn("加高?", default=False),
                    "备注": st.column_config.TextColumn("单行备注")
                },
                key="contract_editor"
            )
            
            # 3. 提交逻辑
            val_note = st.session_state.get('auto_global_note', "")
            f_note_global = st.text_input("合同总备注", value=val_note, placeholder="可选，应用于所有条目")
            
            if st.button("💾 保存所有合同条目", type="primary"):
                # 自动生成合同号
                f_contract = f"HT{datetime.now().strftime('%Y%m%d')}{random.randint(1000, 9999)}"
                
                if edited_df.empty:
                    st.warning("请至少添加一行机型数据")
                else:
                    df_plan = get_factory_plan()
                    new_rows = []
                    
                    # 遍历编辑后的 DataFrame
                    for _, row in edited_df.iterrows():
                        m = row.get('机型')
                        q = int(row.get('数量', 1))
                        is_h = row.get('加高?', False)
                        note_line = str(row.get('备注', ''))
                        
                        if not m: continue # 跳过空机型
                        
                        final_key = f"{m}(加高)" if is_h else m
                        
                        # 合并备注：总备注 + 行备注
                        final_note = f_note_global
                        if note_line:
                            if final_note: final_note += f" | {note_line}"
                            else: final_note = note_line
                        
                        new_rows.append({
                            "合同号": f_contract, "机型": final_key,
                            "排产数量": str(q), "要求交期": str(f_deadline),
                            "状态": "未下单", "备注": final_note,
                            "客户名": f_customer, "代理商": f_agent,
                            "指定批次/来源": ""
                        })
                    
                    if new_rows:
                        df_plan = pd.concat([df_plan, pd.DataFrame(new_rows)], ignore_index=True)
                        save_factory_plan(df_plan)
                        
                        # --- Save File if exists ---
                        if f_file:
                            if not f_customer: f_customer = "Unknown"
                            success, msg = save_contract_file(f_file, f_customer, f_contract, st.session_state.operator_name)
                            if success: st.success(f"📎 附件已上传！")
                            else: st.error(f"❌ 附件保存失败: {msg}")
                            
                        st.success(f"已添加 {len(new_rows)} 条合同记录！"); time.sleep(1); st.rerun()
                    else:
                        st.warning("有效数据为空")
        
        fp = get_factory_plan()
        # 确保新字段存在
        if "客户名" not in fp.columns: fp["客户名"] = ""
        if "代理商" not in fp.columns: fp["代理商"] = ""

        if not fp.empty:
            fp['temp_date'] = pd.to_datetime(fp['要求交期'], errors='coerce').dt.date
            today = datetime.now().date()
            
            tab1, tab2, tab3 = st.tabs(["🔥 紧急提醒 (2周内)", "📅 近期规划 (2月内)", "📋 全景视图"])
            
            def render_plan_table(df_view, key_prefix):
                if df_view.empty:
                    st.info("无相关数据")
                    return
                # 按机型排序
                df_view = df_view.copy()
                df_view['__rank'] = df_view['机型'].apply(get_model_rank)
                df_view = df_view.sort_values(by=['__rank', '要求交期'], ascending=[True, True])
                
                st.dataframe(df_view[["合同号", "客户名", "代理商", "机型", "排产数量", "要求交期", "状态", "备注"]], use_container_width=True, hide_index=True)
                
                c_op1, c_op2 = st.columns([3, 1])
                with c_op1:
                    op_contract = st.selectbox("选择合同号进行操作", df_view['合同号'].unique(), key=f"{key_prefix}_sel")
                with c_op2:
                    # --- [修改] 增加 "关联现有订单" 选项 ---
                    action_opts = ["标记已下单", "标记已完工", "取消计划", "🔗 关联现有订单(核销)"]
                    action = st.radio("动作", action_opts, horizontal=True, key=f"{key_prefix}_act", label_visibility="collapsed")
                    
                    # 如果选择了关联，需要输入订单号
                    link_oid = ""
                    if action == "🔗 关联现有订单(核销)":
                        link_oid = st.text_input("输入已存在的订单号", placeholder="例如: SO-2026...", key=f"{key_prefix}_oid_input")

                    if st.button("执行", key=f"{key_prefix}_btn"):
                        mask = fp['合同号'] == op_contract
                        
                        if action == "标记已下单": 
                            fp.loc[mask, '状态'] = "已下单"
                            st.success("状态已更新！")
                        elif action == "标记已完工": 
                            fp.loc[mask, '状态'] = "已完工"
                            st.success("状态已更新！")
                        elif action == "取消计划": 
                            fp.loc[mask, '状态'] = "已取消"
                            st.success("计划已取消！")
                        elif action == "🔗 关联现有订单(核销)":
                            if not link_oid:
                                st.error("请输入要关联的订单号")
                                st.stop()
                            else:
                                # 检查订单号是否存在
                                all_ords = get_orders()
                                if link_oid not in all_ords['订单号'].values:
                                    st.error(f"订单号 {link_oid} 不存在！请检查拼写。")
                                    st.stop()
                                else:
                                    fp.loc[mask, '状态'] = "已转订单"
                                    fp.loc[mask, '订单号'] = link_oid
                                    st.success(f"已成功将合同 {op_contract} 与订单 {link_oid} 关联！")
                        
                        save_factory_plan(fp)
                        time.sleep(1); st.rerun()

            with tab1:
                deadline_2w = today + timedelta(days=14)
                df_urgent = fp[(fp['temp_date'] <= deadline_2w) & (~fp['状态'].isin(['已完工', '已取消']))]
                render_plan_table(df_urgent, "tab1")
            with tab2:
                deadline_2m = today + timedelta(days=60)
                df_near = fp[(fp['temp_date'] <= deadline_2m) & (~fp['状态'].isin(['已完工', '已取消']))]
                render_plan_table(df_near, "tab2")
            with tab3:
                df_all = fp[~fp['状态'].isin(['已完工', '已取消'])]
                render_plan_table(df_all, "tab3")
        else: st.info("暂无排产计划数据")

# --- ️ 机台信息编辑 ---
elif st.session_state.page == 'machine_edit':
    check_access('MACHINE_EDIT')
    c_back, c_title = st.columns([2, 8])
    with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
    with c_title: st.header("🛠️ 机台信息编辑")

    with st.expander("🔎 筛选条件", expanded=True):
        c_f1, c_f2, c_f3 = st.columns(3)
        with c_f1: f_sn = st.text_input("流水号 (包含)")
        with c_f2: f_order = st.text_input("订单号 (包含)")
        with c_f3: f_date_range = st.date_input("更新时间范围", value=[])
    
    df = get_data()
    edit_df = df[df['状态'] != '已出库'].copy()
    
    if f_sn: edit_df = edit_df[edit_df['流水号'].str.contains(f_sn, case=False, na=False)]
    if f_order: edit_df = edit_df[edit_df['占用订单号'].str.contains(f_order, case=False, na=False)]
    
    if not edit_df.empty:
        # 按机型排序
        edit_df['__rank'] = edit_df['机型'].apply(get_model_rank)
        edit_df = edit_df.sort_values(by=['__rank', '批次号'], ascending=[True, False])
        
        edit_df.insert(0, "✅ 选择", False)
        edited_res = st.data_editor(edit_df[['✅ 选择', '流水号', '机型', '状态', '占用订单号', '机台备注/配置', '更新时间']], hide_index=True, use_container_width=True, key="machine_edit_editor")
        selected_rows = edited_res[edited_res['✅ 选择'] == True]
        
        if not selected_rows.empty:
            st.divider()
            with st.form("batch_edit_form"):
                st.markdown("#### 批量修改信息")
                
                # Layout for new controls
                c_edit1, c_edit2 = st.columns(2)
                with c_edit1:
                    new_model = st.text_input("修改机型 (留空则不修改)")
                with c_edit2:
                    st.caption("快捷选项 (追加到备注)")
                    opt_xs = st.checkbox("XS改X手自一体")
                    opt_cond = st.checkbox("后导电")

                new_note = st.text_area("新的机台备注/配置 (覆盖模式：输入内容+勾选内容)", max_chars=500, help="注意：此处内容将覆盖原有备注。如需保留原备注请勿在此输入且勿勾选，或手动复制原备注。")
                
                if st.form_submit_button("💾 确认修改", type="primary"):
                    sns_val = selected_rows['流水号'].tolist()
                    updates_made = False
                    
                    # 1. 修改机型
                    if new_model.strip():
                        df.loc[df['流水号'].isin(sns_val), '机型'] = new_model.strip()
                        updates_made = True
                    
                    # 2. 修改备注
                    # 逻辑：只要有输入或勾选，就执行更新（覆盖）
                    # 如果全空，则不更新备注（防止误操作清空）
                    if new_note or opt_xs or opt_cond:
                        base_note = new_note if new_note else ""
                        extras = []
                        if opt_xs: extras.append("XS改X手自一体")
                        if opt_cond: extras.append("后导电")
                        
                        final_note = (base_note + " " + " ".join(extras)).strip()
                        
                        df.loc[df['流水号'].isin(sns_val), '机台备注/配置'] = final_note
                        updates_made = True
                    
                    if updates_made:
                        df.loc[df['流水号'].isin(sns_val), '更新时间'] = datetime.now().strftime("%Y-%m-%d %H:%M")
                        save_data(df)
                        st.success(f"已更新 {len(sns_val)} 台机器信息！"); time.sleep(1); st.rerun()
                    else:
                        st.info("未输入任何变更内容")
    else: st.info("无数据")

# --- 📂 机台档案 (Machine Archive) ---
elif st.session_state.page == 'machine_archive':
    check_access('ARCHIVE')
    c_back, c_title = st.columns([2, 8])
    with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
    with c_title: st.header("📂 机台档案 ")

    st.info("💡 管理每台机器的电子档案（照片/文档），文件存储在物理文件夹中。")
    
    # 1. 查询区
    df_all = get_data()
    all_sns = df_all['流水号'].unique().tolist() if not df_all.empty else []
    
    # 支持输入或选择
    if 'archive_sn_search' not in st.session_state: st.session_state.archive_sn_search = ""
    
    col_search, col_info = st.columns([1, 2])
    with col_search:
        # 使用 selectbox 实现带搜索的选择
        selected_sn = st.selectbox("🔍 搜索/选择流水号", [""] + sorted(all_sns, reverse=True), key="archive_sn_select")
    
    if selected_sn:
        # 获取机台信息
        row = df_all[df_all['流水号'] == selected_sn].iloc[0]
        model = row['机型']
        status = row['状态']
        
        with col_info:
            st.markdown(f"### {selected_sn}")
            st.caption(f"机型: {model} | 状态: {status}")
            
        st.divider()
        
        # 准备物理路径
        sn_dir = os.path.join(MACHINE_ARCHIVE_ABS_DIR, selected_sn)
        if not os.path.exists(sn_dir):
            try: os.makedirs(sn_dir, exist_ok=True)
            except: pass
            
        # 2. 展示区 (照片墙)
        # 读取目录下所有图片
        image_files = []
        if os.path.exists(sn_dir):
            all_files = os.listdir(sn_dir)
            image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.gif']
            image_files = [f for f in all_files if os.path.splitext(f)[1].lower() in image_extensions]
            # 按时间倒序
            image_files.sort(key=lambda x: os.path.getmtime(os.path.join(sn_dir, x)), reverse=True)
            
        if image_files:
            st.markdown(f"#### 🖼️ 现有照片 ({len(image_files)} 张)")
            
            # 每行显示 4 张
            cols = st.columns(4)
            for idx, img_name in enumerate(image_files):
                img_path = os.path.join(sn_dir, img_name)
                with cols[idx % 4]:
                    st.image(img_path, caption=img_name, use_container_width=True)
                    # 删除按钮
                    if st.button("🗑️", key=f"del_img_{selected_sn}_{img_name}", help="删除此照片"):
                        try:
                            os.remove(img_path)
                            audit_log("Delete Archive Photo", f"Deleted {img_name} from {selected_sn}")
                            st.rerun()
                        except Exception as e:
                            st.error(f"删除失败: {e}")
        else:
            st.info("暂无照片存档")
            
        st.divider()
        
        # 3. 上传区
        st.markdown("#### 📤 上传机台档案 (必填项)")
        st.info("💡 请输入对应部件编号，并上传照片。系统将自动使用编号重命名文件。")
        
        # 容器布局
        with st.container(border=True):
            # A. 关键部件 (Key Components)
            c_k1, c_k2, c_k3 = st.columns(3)
            
            with c_k1:
                val_wheel = st.text_input("🟢 手轮号 ", key=f"wheel_{selected_sn}")
                file_wheel = st.file_uploader("上传手轮照片", type=['jpg', 'png'], key=f"up_wheel_{selected_sn}")
            
            with c_k2:
                val_motor = st.text_input("🔵 电机号 ", key=f"motor_{selected_sn}")
                file_motor = st.file_uploader("上传电机照片", type=['jpg', 'png'], key=f"up_motor_{selected_sn}")
            
            with c_k3:
                val_board = st.text_input("🟠 板号 ", key=f"board_{selected_sn}")
                file_board = st.file_uploader("上传主板照片", type=['jpg', 'png'], key=f"up_board_{selected_sn}")
        
        st.write("")
        with st.container(border=True):
            # B. 其他照片 (Others)
            st.markdown("##### ⚪ 其他/备注照片")
            c_o1, c_o2 = st.columns([1, 2])
            with c_o1:
                val_other = st.text_input("图片说明 (选填)", placeholder="例如：机身侧面、包装等", key=f"other_txt_{selected_sn}")
            with c_o2:
                files_other = st.file_uploader("上传其他照片 (支持多选)", type=['jpg', 'png'], accept_multiple_files=True, key=f"up_other_{selected_sn}")

        if st.button("💾 保存所有档案照片", type="primary"):
            # 检查必填项：如果上传了图片，则必须有对应的编号
            errors = []
            if file_wheel and not val_wheel.strip(): errors.append("请填写【手轮号】")
            if file_motor and not val_motor.strip(): errors.append("请填写【电机号】")
            if file_board and not val_board.strip(): errors.append("请填写【板号】")
            
            if errors:
                for e in errors: st.error(e)
            else:
                count = 0
                ts_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                def save_arch_file(fileobj, prefix, label_val):
                    if not fileobj: return False
                    ext = os.path.splitext(fileobj.name)[1].lower()
                    if not ext: ext = ".jpg"
                    # Clean filename
                    safe_label = re.sub(r'[\\/*?:"<>|]', "", str(label_val)).strip()
                    if not safe_label: safe_label = "Unamed"
                    
                    # Naming: {Prefix}_{Label}_{Timestamp}.jpg
                    # e.g. Wheel_WH123_20260226.jpg
                    final_name = f"{prefix}_{safe_label}_{ts_str}{ext}"
                    save_p = os.path.join(sn_dir, final_name)
                    
                    try:
                        with open(save_p, "wb") as f: f.write(fileobj.read())
                        return True
                    except: return False

                # 1. Save Wheel
                if save_arch_file(file_wheel, "手轮", val_wheel): count += 1
                # 2. Save Motor
                if save_arch_file(file_motor, "电机", val_motor): count += 1
                # 3. Save Board
                if save_arch_file(file_board, "板号", val_board): count += 1
                
                # 4. Save Others
                if files_other:
                    idx = 1
                    for f_obj in files_other:
                        p_fix = "其他"
                        if val_other.strip(): p_fix = val_other.strip()
                        # Avoid overwrite if multiple files
                        label_comb = f"{idx}"
                        if save_arch_file(f_obj, p_fix, label_comb): count += 1
                        idx += 1
                
                if count > 0:
                    audit_log("Upload Archive", f"Uploaded {count} photos for {selected_sn} (Wheel:{val_wheel}, Motor:{val_motor})")
                    st.success(f"成功归档 {count} 张照片！")
                    time.sleep(1); st.rerun()
                else:
                    st.warning("未检测到待保存的文件")

# ---  销售下单 ---
elif st.session_state.page == 'sales_create':
    check_access('SALES_CREATE')
    c_back, c_title = st.columns([2, 8])
    with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
    with c_title: st.header("📝 销售订单管理")

    # --- 辅助函数：获取全量机型列表 (含0库存) ---
    def get_all_models(df_source):
        all_models = set(CUSTOM_MODEL_ORDER)
        if not df_source.empty:
            all_models.update(df_source['机型'].unique())
        return sorted(list(all_models), key=get_model_rank)

    df_inv = get_data()
    active_inv = df_inv[df_inv['状态'] != '已出库']
    
    # 使用全量模型列表
    available_models = get_all_models(df_inv)
    
    tab_new, tab_import, tab_manage = st.tabs(["➕ 手动下单", "导入已规划合同 ", "订单查询与管理"])
    
    with tab_new:
        # --- Reset Logic ---
        if st.session_state.get("reset_manual_order_flag", False):
            for key in ["mo_cust", "mo_agent", "mo_note", "mo_source", "mo_date", "mo_pack"]:
                if key in st.session_state: del st.session_state[key]
            
            if "manual_order_editor" in st.session_state: del st.session_state["manual_order_editor"]
            st.session_state.manual_order_df = pd.DataFrame(
                [{"机型": available_models[0] if available_models else "", "数量": 1, "加高?": False, "备注": ""}]
            )
            st.session_state["reset_manual_order_flag"] = False

        # --- V7.1 Update: 使用 Data Editor 表格录入 (支持多行/加高) ---
        st.markdown("##### 1. 填写订单详情")
        c_cust, c_agent = st.columns(2)
        with c_cust: inp_customer = st.text_input("客户信息 (Customer)", key="mo_cust")
        with c_agent: inp_agent = st.text_input("代理商 (Agent)", key="mo_agent")

        st.markdown("##### 2. 录入机型")
        st.caption("请在下方表格中添加机型。支持重复添加同一机型（例如一行标准、一行加高）。")

        if 'manual_order_df' not in st.session_state:
            st.session_state.manual_order_df = pd.DataFrame(
                [{"机型": available_models[0] if available_models else "", "数量": 1, "加高?": False, "备注": ""}]
            )

        edited_df = st.data_editor(
            st.session_state.manual_order_df,
            num_rows="dynamic",
            use_container_width=True,
            column_config={
                "机型": st.column_config.SelectboxColumn("机型", options=available_models, required=True),
                "数量": st.column_config.NumberColumn("数量", min_value=1, default=1, required=True),
                "加高?": st.column_config.CheckboxColumn("加高?", default=False),
                "备注": st.column_config.TextColumn("单行备注")
            },
            key="manual_order_editor"
        )
        
        # --- [新增] 深度撞单检测 (Deep Conflict Check) ---
        conflict_found = False
        if inp_customer and not edited_df.empty:
             # 1. 准备用户输入的数据摘要
             user_items = []
             has_user_input = False
             for _, row in edited_df.iterrows():
                 m = row.get('机型')
                 if not m: continue
                 has_user_input = True
                 q = int(row.get('数量', 1))
                 is_h = row.get('加高?', False)
                 final_m = f"{m}(加高)" if is_h else m
                 note = str(row.get('备注', '')).strip()
                 user_items.append({'model': final_m, 'qty': q, 'note': note})
             
             if has_user_input:
                 risk_details = []
                 fp_check = get_factory_plan()
                 # 筛选潜在合同 (状态符合 + 客户名包含)
                 # 修复 re.error: 使用 regex=False 进行纯文本匹配
                 potential_contracts = fp_check[
                     (fp_check['状态'].isin(['已规划', '已审批', '未下单'])) &
                     (fp_check['客户名'].str.contains(inp_customer, na=False, regex=False))
                 ]
                 
                 if not potential_contracts.empty:
                     # 按合同号分组比对
                     for cid, grp in potential_contracts.groupby('合同号'):
                         contract_score = 0
                         match_reasons = []
                         
                         # 构建合同的机型清单
                         c_items = []
                         for _, c_row in grp.iterrows():
                             try: c_q = int(float(c_row['排产数量']))
                             except: c_q = 0
                             c_items.append({'model': c_row['机型'], 'qty': c_q, 'note': str(c_row.get('备注', ''))})
                         
                         # 比对逻辑: 机型匹配+2分, 数量匹配+3分, 备注相似+2分. 阈值>3
                         for u_item in user_items:
                             for c_item in c_items:
                                 if u_item['model'].upper() == c_item['model'].upper():
                                     # 机型匹配
                                     item_score = 2
                                     reason_part = f"{u_item['model']}"
                                     
                                     if u_item['qty'] == c_item['qty']:
                                         item_score += 3
                                         reason_part += f" x{u_item['qty']}(数量一致)"
                                     else:
                                         reason_part += f" (用户:{u_item['qty']} vs 合同:{c_item['qty']})"
                                     
                                     if u_item['note'] and c_item['note'] and (u_item['note'] in c_item['note'] or c_item['note'] in u_item['note']):
                                          item_score += 2
                                          reason_part += " [备注相似]"
                                     
                                     if item_score > 3: # 仅当匹配度较高时才计入风险
                                         contract_score += item_score
                                         match_reasons.append(reason_part)
                         
                         if contract_score >= 5: # 整体风险阈值
                             risk_details.append(f"📄 **合同 {cid}**: 包含 {', '.join(set(match_reasons))}")

                 # 2. 检测现有订单 (Sales Orders)
                 existing_orders = get_orders()
                 if not existing_orders.empty:
                    # 查找同客户名的订单
                    # 修复 re.error: 使用 regex=False 进行纯文本匹配
                    potential_orders = existing_orders[
                         existing_orders['客户名'].str.contains(inp_customer, na=False, regex=False) | 
                         existing_orders['客户名'].apply(lambda x: inp_customer in str(x))
                    ]
                    
                    if not potential_orders.empty:
                         # user_items 转为 dict: {model_key: qty}
                         user_reqs_check = {}
                         for u in user_items:
                             user_reqs_check[u['model']] = user_reqs_check.get(u['model'], 0) + u['qty']
                         
                         for _, ord_row in potential_orders.iterrows():
                             # 解析订单需求
                             ord_reqs = parse_requirements(ord_row['需求机型'], ord_row['需求数量'])
                             
                             if user_reqs_check == ord_reqs:
                                 o_time = ord_row.get('下单时间', '未知时间')
                                 risk_details.append(f"📦 **现有订单 {ord_row['订单号']}**: 内容完全一致！(下单时间: {o_time})")

                 if risk_details:
                     st.error(
                         f"🚨 **深度撞单预警**：检测到以下合同/订单与当前录入高度雷同！\n\n" + 
                         "\n".join(risk_details) + 
                         "\n\n👉 请务必检查！如需强制下单，请勾选下方的【确认】框。"
                     )
                     conflict_found = True

        st.markdown("---")
        c2_1, c2_2 = st.columns([3, 1])
        with c2_1:
            inp_note_global = st.text_input("订单总备注", key="mo_note")
            inp_delivery_date = st.date_input("发货时间 (选填)", value=None, key="mo_date")
        with c2_2:
            st.write(""); st.write("")
            need_pack = st.checkbox("需要包装", key="mo_pack")
        
        inp_source = st.text_input("指定批次/来源 (初始备注)", placeholder="如：优先现货", key="mo_source")
        
        confirm_force = False
        if conflict_found:
             confirm_force = st.checkbox("⚠️ 我确认这不是重复下单 (I confirm this is NOT a duplicate)", key="force_submit_duplicate")

        if st.button("✅ 生成订单", type="primary", use_container_width=True):
            if conflict_found and not confirm_force:
                 st.error("❌ 操作已拦截：检测到重复下单风险。请勾选上方的确认框以继续。")
                 st.stop()
            
            if not inp_customer: 
                st.error("请输入客户信息")
            elif edited_df.empty:
                st.warning("请至少添加一行机型数据")
            else:
                model_qty_map = {}
                combined_notes = []
                if inp_note_global: combined_notes.append(inp_note_global)
                
                has_valid_row = False
                for _, row in edited_df.iterrows():
                    m = row.get('机型')
                    q = int(row.get('数量', 1))
                    is_h = row.get('加高?', False)
                    r_note = str(row.get('备注', '')).strip()
                    
                    if not m: continue
                    has_valid_row = True
                    
                    final_key = f"{m}(加高)" if is_h else m
                    model_qty_map[final_key] = model_qty_map.get(final_key, 0) + q
                    
                    if r_note:
                        combined_notes.append(f"[{final_key}: {r_note}]")
                
                if not has_valid_row:
                    st.error("有效机型数据为空")
                else:
                    final_note_str = " ".join(combined_notes)
                    pack_opt = "需要包装" if need_pack else "不包装"
                    delivery_str = inp_delivery_date.strftime("%Y-%m-%d") if inp_delivery_date else ""
                    
                    oid = create_sales_order(inp_customer, inp_agent, model_qty_map, final_note_str, pack_opt, delivery_str, inp_source)
                    
                    # Set flag to reset on next run
                    st.session_state["reset_manual_order_flag"] = True
                    
                    st.success(f"订单已生成: {oid}"); time.sleep(1); st.rerun()

    with tab_import:
        st.subheader("📥 导入已规划合同 (Import Planned Contracts)")
        fp_df = get_factory_plan()
        if "客户名" not in fp_df.columns: fp_df["客户名"] = ""
        
        # 筛选 '已规划' 的合同
        planned_contracts = fp_df[fp_df['状态'] == '已规划'].copy()
        
        if planned_contracts.empty:
            st.info("暂无已规划的合同")
        else:
            # 显示列表
            st.dataframe(planned_contracts[["合同号", "客户名", "机型", "排产数量", "要求交期", "备注"]], use_container_width=True, hide_index=True)
            
            st.divider()
            
            # --- 新逻辑：按合同号聚合 ---
            # 1. 获取唯一合同号
            unique_contracts = planned_contracts['合同号'].unique()
            
            # 2. 构建选项
            contract_opts = []
            contract_map = {} # label -> contract_id
            
            for cid in unique_contracts:
                c_rows = planned_contracts[planned_contracts['合同号'] == cid]
                cust = c_rows.iloc[0]['客户名']
                # 汇总机型
                models_list = c_rows['机型'].unique()
                total_qty = c_rows['排产数量'].astype(float).sum()
                
                label = f"{cid} | {cust} | 共 {int(total_qty)} 台 ({len(models_list)} 款机型)"
                contract_opts.append(label)
                contract_map[label] = cid

            # --- V7.3 Update: 支持多选合并 ---
            sel_strs = st.multiselect("选择要转换/合并的合同 (支持多选)", contract_opts)
            
            if sel_strs:
                sel_cids = [contract_map[s] for s in sel_strs]
                
                # 获取所有选中合同的行
                all_target_rows = planned_contracts[planned_contracts['合同号'].isin(sel_cids)]
                
                # 基础信息取第一个合同的作为默认值
                first_row = all_target_rows.iloc[0]
                
                # 检查客户是否一致
                unique_customers = all_target_rows['客户名'].unique()
                if len(unique_customers) > 1:
                    st.warning(f"⚠️ 注意：您选择了不同客户的合同进行合并 ({', '.join(unique_customers)})，请确认是否正确。")
                
                st.markdown("#### 📝 确认合并订单信息 (可修改)")
                with st.form("confirm_planned_order"):
                    c1, c2 = st.columns(2)
                    with c1: 
                        new_cust = st.text_input("客户名", value=first_row.get('客户名', ''))
                        new_agent = st.text_input("代理商", value=first_row.get('代理商', ''))
                    with c2:
                        new_delivery = st.text_input("发货时间/交期", value=first_row.get('要求交期', ''))
                        new_pack = st.checkbox("需要包装", value=False)
                    
                    st.write("**包含机型及数量 (合并汇总):**")
                    
                    # --- 准备合并数据 ---
                    model_lines = []
                    # 这里的 target_rows 包含所有选中的合同行
                    for idx, row in all_target_rows.iterrows():
                        raw_model = row['机型']
                        is_h = "(加高)" in raw_model
                        base_m = raw_model.replace("(加高)", "")
                        
                        # 查找是否已存在该机型 (为了UI合并显示，避免列表太长，但为了保留原始备注，最好还是列出来?)
                        # 用户需求是合并出货。
                        # 如果我们把相同机型合并成一行，备注怎么处理？
                        # 方案：不合并行，罗列所有行，用户确认总数。或者：按机型合并，备注拼接。
                        # 这里选择：不合并行，让用户看到每一笔来源，这样更清晰。
                        # 但为了方便，可以在表格里加一列 "来源合同"
                        
                        model_lines.append({
                            "来源合同": row['合同号'],
                            "机型": base_m,
                            "加高?": is_h,
                            "数量": int(float(row['排产数量'])),
                            "原备注": row.get('备注', ''),
                            "__idx": idx
                        })
                    
                    # 使用 data_editor
                    df_models_confirm = pd.DataFrame(model_lines)
                    edited_models = st.data_editor(
                        df_models_confirm[['来源合同', '机型', '加高?', '数量', '原备注']],
                        key="editor_contract_models",
                        use_container_width=True,
                        disabled=["来源合同", "机型", "原备注"],
                        column_config={
                            "加高?": st.column_config.CheckboxColumn(
                                "加高?",
                                help="勾选后将自动添加 (加高) 后缀",
                                default=False,
                            ),
                            "数量": st.column_config.NumberColumn("数量", min_value=1)
                        }
                    )
                    
                    # 构建默认备注（包含各机型的备注）
                    default_note = ""
                    # 简单的去重合并备注
                    seen_notes = set()
                    for idx, row in all_target_rows.iterrows():
                         r_note = str(row.get('备注', '')).strip()
                         if r_note and r_note not in seen_notes:
                             cid_prefix = f"[{row['合同号']}] " if len(sel_cids) > 1 else ""
                             default_note += f"{cid_prefix}{r_note} "
                             seen_notes.add(r_note)

                    new_note = st.text_area("订单总备注", value=default_note.strip())
                    
                    if st.form_submit_button("🚀 确认生成合并订单 (Confirm Merge)", type="primary"):
                        pack_opt = "需要包装" if new_pack else "不包装"
                        
                        # 1. 收集机型数据 (合并同类项)
                        final_model_data = {}
                        
                        for _, m_row in edited_models.iterrows():
                            m_name = m_row['机型']
                            is_h = m_row['加高?']
                            final_name = f"{m_name}(加高)" if is_h else m_name
                            
                            m_qty = int(m_row['数量'])
                            final_model_data[final_name] = final_model_data.get(final_name, 0) + m_qty
                        
                        # 2. 合并 Source Batch (Plan String)
                        # 必须解析 -> 合并 -> 序列化，防止覆盖
                        merged_plan_map = {} # {Model: {Batch: Qty}}
                        
                        for idx, row in all_target_rows.iterrows():
                            p_str = str(row.get('指定批次/来源', ''))
                            if not p_str: continue
                            
                            # 解析单个 plan string (e.g. "ModelA:{...}; ModelB:{...}")
                            parts = p_str.split(";")
                            for part in parts:
                                if ":" in part:
                                    try:
                                        m_key, content = part.split(":", 1)
                                        m_key = m_key.strip()
                                        alloc_dict = ast.literal_eval(content.strip())
                                        
                                        if m_key not in merged_plan_map:
                                            merged_plan_map[m_key] = {}
                                        
                                        # Merge alloc_dict
                                        for batch, qty in alloc_dict.items():
                                            merged_plan_map[m_key][batch] = merged_plan_map[m_key].get(batch, 0) + int(qty)
                                    except: pass
                        
                        # 序列化
                        all_plans_final = []
                        for m_key, alloc_data in merged_plan_map.items():
                            all_plans_final.append(f"{m_key}:{str(alloc_data)}")
                            
                        combined_plan_str = "; ".join(all_plans_final)
                        
                        # 3. 创建订单
                        new_oid = create_sales_order(
                            customer=new_cust,
                            agent=new_agent,
                            model_data=final_model_data,
                            note=new_note,
                            pack_option=pack_opt,
                            delivery_time=new_delivery,
                            source_batch=combined_plan_str
                        )
                        
                        # 4. 更新所有合同状态
                        fp_df.loc[fp_df['合同号'].isin(sel_cids), '状态'] = '已转订单'
                        fp_df.loc[fp_df['合同号'].isin(sel_cids), '订单号'] = new_oid
                        save_factory_plan(fp_df)
                        
                        st.success(f"已生成合并订单: {new_oid}！包含 {len(sel_cids)} 份合同。"); time.sleep(1); st.rerun()

    with tab_manage:
        # 1. 获取数据
        q_orders = get_orders()
        df_inv = get_data()
        
        # 2. 计算订单完成度
        if not df_inv.empty:
            shipped_stats = df_inv[df_inv['状态'] == '已出库'].groupby('占用订单号').size().to_dict()
        else: shipped_stats = {}
        
        def check_order_status(oid, req_qty_str, db_status):
            # 优先判断数据库标记的状态
            if str(db_status) == 'deleted': return 'deleted'
            
            if not oid: return "unknown"
            s_qty = shipped_stats.get(oid, 0)
            try: r_qty = int(float(req_qty_str))
            except: r_qty = 999999
            
            if s_qty >= r_qty and r_qty > 0: return "completed"
            return "active"

        # 3. 筛选器
        st.markdown("#### 🔍 订单查询与管理")
        
        # 4. 过滤数据
        if not q_orders.empty:
            # 添加状态列辅助筛选
            q_orders['__status'] = q_orders.apply(lambda r: check_order_status(r['订单号'], r['需求数量'], r.get('status', '')), axis=1)
            
            # --- V7.2 新增：按月份筛选 ---
            # 确保下单时间为 datetime
            q_orders['下单时间_dt'] = pd.to_datetime(q_orders['下单时间'], errors='coerce')
            q_orders['month_str'] = q_orders['下单时间_dt'].dt.strftime('%Y-%m')
            
            available_months = sorted(q_orders['month_str'].dropna().unique().tolist(), reverse=True)
            month_opts = ["全部"] + available_months
            
            c_f1, c_f2 = st.columns([3, 1])
            with c_f1:
                filter_status = st.radio("订单状态筛选", ["进行中 (Active)", "往期/已完结 (Completed)", "已删除 (Deleted)"], horizontal=True)
            with c_f2:
                sel_month = st.selectbox("📅 按下单月份筛选", month_opts)
            
            target_status = "active"
            if "已完结" in filter_status: target_status = "completed"
            elif "已删除" in filter_status: target_status = "deleted"
            
            # 构建过滤掩码
            mask = (q_orders['__status'] == target_status)
            if sel_month != "全部":
                mask = mask & (q_orders['month_str'] == sel_month)
            
            view_df = q_orders[mask].copy()
            
            # 搜索框
            search_txt = st.text_input("搜索订单 (订单号/客户/代理)", key="manage_search")
            if search_txt:
                s = search_txt.lower()
                view_df = view_df[
                    view_df['订单号'].str.lower().str.contains(s, na=False) |
                    view_df['客户名'].str.lower().str.contains(s, na=False) |
                    view_df['代理商'].str.lower().str.contains(s, na=False)
                ]
            
            # 5. 显示与编辑
            if view_df.empty:
                st.info("无相关订单数据")
            else:
                # 倒序显示
                view_df = view_df.iloc[::-1]
                
                # 如果是“进行中”订单，允许编辑
                if target_status == "active":
                    st.caption(f"共找到 {len(view_df)} 个进行中订单。支持直接编辑【备注】和【发货时间】，勾选后可删除。")
                elif target_status == "deleted":
                    st.caption(f"共找到 {len(view_df)} 个已删除订单。")
                else:
                    st.caption(f"共找到 {len(view_df)} 个已完结订单。")
                    
                # 预处理：确保 '发货时间' 是 datetime 类型
                view_df['发货时间'] = pd.to_datetime(view_df['发货时间'], errors='coerce').dt.date
                
                # 需要显示的列
                disp_cols = ["订单号", "客户名", "代理商", "需求机型", "需求数量", "发货时间", "备注", "下单时间"]
                
                # 如果是已删除，显示删除原因
                if target_status == "deleted":
                    if "delete_reason" not in view_df.columns: view_df['delete_reason'] = ""
                    disp_cols.append("delete_reason")
                
                # 允许删除的操作列 (仅 Active/Completed)
                if target_status != "deleted":
                    view_df.insert(0, "✅", False)
                    final_cols = ["✅"] + disp_cols
                else:
                    final_cols = disp_cols
                
                # 配置 column config
                column_config = {
                    "订单号": st.column_config.TextColumn(disabled=True),
                    "客户名": st.column_config.TextColumn(disabled=True),
                    "代理商": st.column_config.TextColumn(disabled=True),
                    "需求机型": st.column_config.TextColumn(disabled=True),
                    "需求数量": st.column_config.TextColumn(disabled=True),
                    "下单时间": st.column_config.TextColumn(disabled=True),
                    "发货时间": st.column_config.DateColumn("发货时间", format="YYYY-MM-DD"),
                    "备注": st.column_config.TextColumn("备注 (可编辑)"),
                    "delete_reason": st.column_config.TextColumn("删除原因", disabled=True),
                    "✅": st.column_config.CheckboxColumn("选择", default=False)
                }
                
                # data_editor
                edited_df = st.data_editor(
                    view_df[final_cols],
                    hide_index=True,
                    use_container_width=True,
                    column_config=column_config,
                    key="order_manage_editor"
                )
                
                # --- 保存修改逻辑 (仅针对备注/时间) ---
                if target_status == "active": # 只有 Active 状态通常允许改这些，Completed 也可以但这里限制一下？原逻辑是 Active 允许
                    if st.button("💾 保存信息修改", type="primary"):
                        changed_count = 0
                        for idx, row in edited_df.iterrows():
                            oid = row['订单号']
                            new_note = row['备注']
                            new_date = row['发货时间'].strftime("%Y-%m-%d") if row['发货时间'] else ""
                            
                            mask = q_orders['订单号'] == oid
                            if not q_orders[mask].empty:
                                org_row = q_orders[mask].iloc[0]
                                org_note = org_row['备注']
                                org_date = org_row['发货时间']
                                
                                if str(new_note) != str(org_note) or str(new_date) != str(org_date):
                                    q_orders.loc[mask, '备注'] = new_note
                                    q_orders.loc[mask, '发货时间'] = new_date
                                    changed_count += 1
                        
                        if changed_count > 0:
                            # 清理临时列
                            for c in ['__status', 'month_str', '下单时间_dt']: 
                                if c in q_orders.columns: del q_orders[c]
                            save_orders(q_orders)
                            st.success(f"已更新 {changed_count} 条订单信息！"); time.sleep(1); st.rerun()
                        else:
                            st.info("未检测到修改")

                # --- 删除逻辑 ---
                if target_status != "deleted":
                    to_delete = edited_df[edited_df['✅'] == True]
                    if not to_delete.empty:
                        st.divider()
                        st.markdown(f"#### 🗑️ 删除订单操作 (选中 {len(to_delete)} 个)")
                        with st.form("delete_order_form"):
                            del_reason = st.text_input("请输入删除原因 (必填):", placeholder="例如：客户取消、重复下单等")
                            if st.form_submit_button("⚠️ 确认删除", type="secondary"):
                                if not del_reason.strip():
                                    st.error("❌ 必须填写删除原因才能删除！")
                                else:
                                    oids_to_del = to_delete['订单号'].tolist()
                                    # Update Status
                                    q_orders.loc[q_orders['订单号'].isin(oids_to_del), 'status'] = 'deleted'
                                    q_orders.loc[q_orders['订单号'].isin(oids_to_del), 'delete_reason'] = del_reason
                                    
                                    # Cleanup and Save
                                    for c in ['__status', 'month_str', '下单时间_dt']: 
                                        if c in q_orders.columns: del q_orders[c]
                                    save_orders(q_orders)
                                    st.success(f"已删除 {len(oids_to_del)} 个订单！"); time.sleep(1); st.rerun()

        else:
            st.info("暂无订单记录")

# --- 📦 订单配货 (升级版：显示老板指示) ---
elif st.session_state.page == 'sales_alloc':
    check_access('SALES_ALLOC')
    c_back, c_title = st.columns([2, 8])
    with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
    with c_title: st.header("📦 订单智能配货")

    orders = get_orders()
    inventory = get_data()
    
    if not inventory.empty:
        shipped_total_stats = inventory[inventory['状态'] == '已出库'].groupby('占用订单号').size().to_dict()
    else: shipped_total_stats = {}
    
    active_orders = orders.iloc[::-1]
    
    # 过滤掉已删除的订单
    if 'status' in active_orders.columns:
        active_orders = active_orders[active_orders['status'] != 'deleted']
    
    if active_orders.empty: st.info("暂无订单。")
    else:
        # 增加筛选：仅显示有老板指示的订单
        filter_has_plan = st.checkbox("🔍 仅显示有老板指示的订单", value=False)

        for idx, row in active_orders.iterrows():
            oid = row['订单号']
            customer = row['客户名']; agent = row['代理商']; note = str(row['备注'])
            
            # 获取发货时间
            raw_date = row.get('发货时间', '')
            delivery_date = str(raw_date) if pd.notna(raw_date) and str(raw_date).strip() != '' else "未指定"
            
            # --- 显示老板的规划指示 ---
            # V7.2: 优先从 CSV 读取规划记录，解决覆盖更新问题
            source_plan = ""
            try:
                plan_records = get_planning_records()
                order_plans = plan_records[plan_records['order_id'] == oid]
                
                if not order_plans.empty:
                    combined_plans = []
                    for _, pr in order_plans.iterrows():
                        p_info = str(pr['plan_info'])
                        p_model = str(pr['model'])
                        combined_plans.append(f"{p_model}:{p_info}")
                    source_plan = "; ".join(combined_plans)
                else:
                    source_plan = row.get('指定批次/来源', '')
            except Exception as e:
                print(f"Error loading plan from CSV: {e}")
                source_plan = row.get('指定批次/来源', '')

            plan_html = ""
            has_valid_plan = False
            
            if source_plan:
                # 智能过滤：去除空指示 (例如 "Model: {}")
                raw_items = str(source_plan).split(";")
                valid_items = []
                for item in raw_items:
                    s_item = item.strip()
                    if not s_item: continue
                    
                    # 检查是否实质为空 (以 {} 或 [] 结尾)
                    # 去除所有空格后检查
                    clean_check = s_item.replace(" ", "")
                    if clean_check.endswith("{}") or clean_check.endswith("[]"):
                        continue
                    # 简单的 Key: 格式也忽略
                    if clean_check.endswith(":"):
                        continue
                        
                    valid_items.append(s_item)
                
                if valid_items:
                    has_valid_plan = True
                    # 重新组合并美化
                    cleaned_source = "; ".join(valid_items)
                    display_plan = cleaned_source.replace(";", "<br>").replace("{", " [").replace("}", "] ").replace("'", "")
                    plan_html = f"<div style='background:#FFF8DC; color:#8B4500; padding:5px; border-radius:4px; font-size:14px;'><b>👑 老板指示:</b><br>{display_plan}</div>"
            
            # 筛选逻辑
            if filter_has_plan and not has_valid_plan:
                continue

            requirements = parse_requirements(row['需求机型'], row['需求数量'])
            total_req_qty = sum(requirements.values())
            
            current_alloc_df = inventory[inventory['占用订单号'] == oid]
            current_total_filled = len(current_alloc_df)
            shipped_count = shipped_total_stats.get(oid, 0)
            
            if shipped_count >= total_req_qty and total_req_qty > 0: continue
            
            # 使用单行 HTML 避免 Markdown 解析缩进问题
            st.markdown(f"""<div class="order-card" style="border:1px solid #ddd; padding:10px; border-radius:5px; margin-bottom:10px;"><h4>📜 {oid} | {customer} (代理: {agent})</h4><p><b>📅 发货时间: {delivery_date}</b></p>{plan_html}<p><b>进度: {current_total_filled} / {total_req_qty}</b> (已发: {shipped_count})</p><p style="color:gray; font-size:14px;">📝 {note}</p></div>""", unsafe_allow_html=True)
            
            for model_key, target_qty in requirements.items():
                # --- V7.1 核心解析 ---
                is_high_req = "(加高)" in model_key
                real_model = model_key.replace("(加高)", "")
                
                # 显示给用户的标题
                display_name = f"{real_model} (加高)" if is_high_req else real_model
                
                # 获取该特定需求已分配的数量
                alloc_mask = (current_alloc_df['机型'] == real_model)
                if is_high_req:
                    alloc_mask = alloc_mask & (current_alloc_df['机台备注/配置'].str.contains("加高", na=False))
                else:
                    alloc_mask = alloc_mask & (~current_alloc_df['机台备注/配置'].str.contains("加高", na=False))
                
                allocated_for_model = len(current_alloc_df[alloc_mask])
                
                remaining = target_qty - allocated_for_model
                status_icon = "✅" if remaining <= 0 else "⏳"
                
                with st.expander(f"{status_icon} 机型: {display_name} | 缺: {max(0, remaining)}", expanded=(remaining > 0)):
                    # --- 显示机型备注 ---
                    model_remark = ""
                    try:
                        # 尝试从订单备注中提取 "Model:Remark"
                        # 优先匹配 model_key (原始键), 然后尝试 display_name
                        keys_to_try = [model_key, display_name]
                        for k in keys_to_try:
                            if not k: continue
                            safe_k = re.escape(k)
                            # 匹配: (开头或空格)Key:(内容)(直到 空格+Word+: 或 结尾)
                            pattern = rf"(?:^|\s){safe_k}:(.*?)(?=\s[\w\(\)\-]+\:|{'$'})"
                            match = re.search(pattern, note)
                            if match:
                                model_remark = match.group(1).strip()
                                break
                    except: pass
                    
                    if model_remark:
                        st.info(f"📝 **备注:** {model_remark}")
                    
                    # --- V7.2 Update: 扩展加高定义 ---
                    # 如果订单备注/机台备注中包含 "加高"，也视为加高需求
                    if "加高" in model_remark:
                        is_high_req = True
                        st.caption("ℹ️ 检测到备注包含“加高”，自动匹配加高库存。")

                    if remaining > 0:
                        # --- V7.1 库存过滤核心 ---
                        mask = (
                            (inventory['机型'] == real_model) & 
                            (inventory['状态'].isin(['待入库', '库存中'])) & 
                            (inventory['占用订单号'] == "")
                        )
                        # --- V7.3 统一加高判断逻辑 ---
                        # 定义：任何字段包含 "加高" 即视为加高库存
                        is_stock_high = (
                            inventory['机型'].str.contains("加高", na=False) |
                            inventory['机台备注/配置'].str.contains("加高", na=False) |
                            inventory['订单备注'].str.contains("加高", na=False)
                        )

                        if is_high_req:
                            mask = mask & is_stock_high
                            st.info("🎯 已自动过滤为：【加高】机器")
                        else:
                            # 需求是标准：排除所有“加高”特征的
                            mask = mask & (~is_stock_high)
                        
                        available_stock = inventory[mask]
                        
                        if available_stock.empty:
                            st.warning(f"⚠️ {display_name} 暂无可用库存")
                        else:
                            # 增加批次筛选功能，方便员工按老板指示找货
                            batches_avail = available_stock['批次号'].unique()
                            c_filter1, c_filter2 = st.columns([1, 2])
                            with c_filter1:
                                filter_b = st.selectbox("按批次筛选 (参考老板指示)", ["全部"] + list(batches_avail), key=f"filter_b_{oid}_{model_key}")
                            
                            filtered_stock = available_stock
                            if filter_b != "全部":
                                filtered_stock = available_stock[available_stock['批次号'] == filter_b]

                            st.markdown(f"**勾选配货 (需 {remaining} 台):**")
                            select_pool = filtered_stock[['批次号', '流水号', '状态', '机台备注/配置', '订单备注']].reset_index(drop=True)
                            # 虽然是单机型，但也按批次排序
                            select_pool = select_pool.sort_values(by=['批次号'], ascending=False)
                            
                            select_pool.insert(0, "✅ 选择", False)
                            
                            key_alloc = f"alloc_{oid}_{model_key}"
                            edited_pool = st.data_editor(select_pool, key=key_alloc, hide_index=True, use_container_width=True, height=200)
                            
                            selected_rows = edited_pool[edited_pool['✅ 选择'] == True]
                            if not selected_rows.empty:
                                if st.button(f"🚀 确认分配 {len(selected_rows)} 台", key=f"btn_go_{key_alloc}", type="primary"):
                                    allocate_inventory(oid, customer, agent, selected_rows['流水号'].tolist())
                                    st.success("成功！"); time.sleep(0.5); st.rerun()

            # 撤回逻辑 (保持原样)
            if current_total_filled > 0:
                with st.expander(f"🔄 配货撤回 ({current_total_filled})", expanded=False):
                    revertable = inventory[(inventory['占用订单号'] == oid) & (inventory['状态'] != '已出库')].copy()
                    if not revertable.empty:
                        revertable.insert(0, "❌", False)
                        res_rev = st.data_editor(revertable[['❌', '批次号', '流水号', '机型']], key=f"rev_{oid}", hide_index=True)
                        to_rev = res_rev[res_rev['❌'] == True]
                        if not to_rev.empty and st.button("确认撤回", key=f"btn_rev_{oid}"):
                            revert_to_inbound(to_rev['流水号'].tolist()); st.rerun()
    
    st.divider()
    render_module_logs(["配货锁定", "自动入库", "撤回"])



# --- 🚛 发货复核 ---
elif st.session_state.page == 'ship_confirm':
    check_access('SHIP_CONFIRM')
    c_back, c_title = st.columns([2, 8])
    with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
    with c_title: st.header("🚛 发货复核")
    
    df = get_data()
    pending = df[df['状态'] == '待发货']
    
    st.metric("待发货总数", len(pending))
    
    if pending.empty: st.success("无任务")
    else:
        # 关联最新的订单备注
        orders_df = get_orders()
        if not orders_df.empty:
            # map order note
            note_map = orders_df.set_index('订单号')['备注'].to_dict()
            # map delivery time
            date_map = orders_df.set_index('订单号')['发货时间'].to_dict()
            
            # 更新订单备注 (如果订单中有备注则使用订单的，否则保留原样)
            if '订单备注' not in pending.columns: pending['订单备注'] = ""
            # 使用 map 更新，注意要处理 NaN
            mapped_notes = pending['占用订单号'].map(note_map)
            pending['订单备注'] = mapped_notes.fillna(pending['订单备注'])
            
            # 更新发货时间
            raw_dates = pending['占用订单号'].map(date_map)
            pending['发货时间'] = pd.to_datetime(raw_dates, errors='coerce').dt.date
        else:
            if '发货时间' not in pending.columns: pending['发货时间'] = None

        # 按机型排序
        pending = pending.copy()
        pending['__rank'] = pending['机型'].apply(get_model_rank)
        pending = pending.sort_values(by=['__rank', '流水号'], ascending=[True, False])
        
        pending.insert(0, "✅", False)
        
        # 显示列包括 订单备注 和 机台备注/配置
        cols_to_show = ['✅', '发货时间', '占用订单号', '客户', '机型', '流水号', '订单备注', '机台备注/配置']
        # 确保列存在
        for c in cols_to_show:
            if c not in pending.columns: pending[c] = ""
            
        res = st.data_editor(
            pending[cols_to_show], 
            hide_index=True, 
            use_container_width=True,
            column_config={
                "发货时间": st.column_config.DateColumn("发货时间", format="YYYY-MM-DD", width="small"),
                "订单备注": st.column_config.TextColumn("订单备注", width="medium"),
                "机台备注/配置": st.column_config.TextColumn("机台备注", width="medium")
            }
        )
        to_act = res[res['✅'] == True]
        
        if not to_act.empty:
            # --- 📸 选中机台照片预览 (Added) ---
            st.divider()
            st.markdown("### 📸 选中机台照片预览")
            for _, row in to_act.iterrows():
                sn = row['流水号']
                model = row['机型']
                with st.expander(f"📦 {sn} - {model}", expanded=True):
                    render_archive_preview(sn)
            st.divider()
            # -----------------------------------

            c_op1, c_op2 = st.columns([1, 1])
            with c_op1:
                if st.button(f"🚚 正式发货 {len(to_act)} 台", type="primary", use_container_width=True):
                    sns = to_act['流水号'].tolist()
                    df.loc[df['流水号'].isin(sns), '状态'] = '已出库'
                    df.loc[df['流水号'].isin(sns), '更新时间'] = datetime.now().strftime("%Y-%m-%d %H:%M")
                    save_data(df)
                    archive_shipped_data(df[df['流水号'].isin(sns)])
                    append_log("正式发货", sns)
                    st.success("发货完成"); time.sleep(1); st.rerun()
            
            with c_op2:
                # 发货撤回功能：撤回为待入库
                if st.button(f"🔄 撤回 {len(to_act)} 台 (退回待入库)", type="secondary", use_container_width=True):
                    revert_to_inbound(to_act['流水号'].tolist(), reason="正式发货撤回")
                    st.success("已撤回为待入库状态！"); time.sleep(1); st.rerun()

    st.divider()
    render_module_logs(["正式发货", "正式发货撤回", "已出库"])

# --- 📥 入库 (V5.5) ---
elif st.session_state.page == 'inbound':
    check_access('INBOUND')
    c_back, c_title = st.columns([2, 8])
    with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
    with c_title: st.header("📥 入库作业")
    
    # 权限检查
    check_prod_admin_permission()

    tab_machine, tab_import = st.tabs(["🏭 机台入库 (Machine Inbound)", "📋 跟踪单导入 (Tracking Import)"])
    
    # --- 模块一：机台入库 ---
    with tab_machine:
        render_machine_inbound_module()
        
    # --- 模块二：跟踪单导入 ---
    with tab_import:
        render_tracking_import_module()
        
        st.divider()
        with st.expander("⚡ 辅助功能：自动生成流水号 (Auto Generate)", expanded=False):
            st.caption("用于生成测试数据或无跟踪单的情况，生成后将写入 PLAN_IMPORT。")
            c1, c2, c3, c4 = st.columns(4)
            with c1: inp_batch = st.text_input("批次号", key="auto_batch")
            with c2:
                df_for_model = get_data()
                all_known_models = set(CUSTOM_MODEL_ORDER)
                if not df_for_model.empty:
                    all_known_models.update(df_for_model['机型'].unique())
                models = sorted(list(all_known_models), key=get_model_rank)
                
                inp_model = st.selectbox("机型", options=models + ["其它(手输)"] if models else ["其它(手输)"], key="auto_model_sel")
                if inp_model == "其它(手输)":
                    final_model = st.text_input("请输入机型名称", key="auto_model_txt")
                else: final_model = inp_model
            with c3: inp_qty = st.number_input("数量", min_value=1, value=1, key="auto_qty")
            with c4: inp_date = st.date_input("预计入库", value=datetime.now().date(), key="auto_date")
            inp_note = st.text_area("机台备注", key="auto_note")
            
            confirm_gen = st.checkbox("我确认上述信息无误", key="auto_confirm")
            if st.button("✅ 生成并保存到 PLAN_IMPORT", type="primary", disabled=not confirm_gen, key="auto_btn"):
                code, msg = generate_auto_inbound(inp_batch, final_model, inp_qty, inp_date, inp_note)
                if code == 1:
                    st.success(msg)
                    time.sleep(1); st.rerun()
                else: st.error(msg)

# --- 🔍 查询 (保持原版逻辑) ---
elif st.session_state.page == 'query':
    c_back, c_title = st.columns([2, 8])
    with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
    with c_title: st.header("📊 库存全景查询")
    
    # --- Dashboard Mode ---
    df = get_data()
    # 始终包含 '库存中' 和 '待入库'
    valid_df = df[df['状态'].isin(['库存中', '待入库'])].copy()
    
    # --- 1. 机型筛选 (Optional) ---
    all_known_models = set(CUSTOM_MODEL_ORDER)
    if not df.empty: all_known_models.update(df['机型'].unique())
    unique_models = sorted(list(all_known_models), key=get_model_rank)
    
    # 顶部只留筛选
    c_q1, c_q2 = st.columns([3, 1])
    with c_q1:
        selected_models_query = st.multiselect("筛选机型", unique_models)
    with c_q2:
        st.write(""); st.write("")
        show_high_only = st.checkbox("仅显示加高 (High Only)")
    
    # 根据筛选过滤 valid_df 用于计算比例和图表
    if selected_models_query:
        display_df = valid_df[valid_df['机型'].isin(selected_models_query)]
        display_models = selected_models_query
    else:
        display_df = valid_df
        display_models = unique_models
        
    if show_high_only:
        display_df = display_df[
            display_df['机型'].str.contains("加高", na=False) |
            display_df['机台备注/配置'].str.contains("加高", na=False) |
            display_df['订单备注'].str.contains("加高", na=False)
        ]

    # --- 2. 库存比例看板 (直接显示) ---
    st.markdown("### 🧮 库存比例 (看板)")
    cols = st.columns(4)
    
    for idx, (label, (mA, mB)) in enumerate(PRESET_RATIOS.items()):
        with cols[idx % 4]:
            # 计算逻辑：使用 display_df (当前视图下的数据)
            cnt_a = len(display_df[display_df['机型'].isin(mA)])
            cnt_b = len(display_df[display_df['机型'].isin(mB)])
            pct = (cnt_a / cnt_b * 100) if cnt_b > 0 else 0.0
            
            # 使用 metric 直接显示
            st.metric(label=label, value=f"{pct:.1f}%", delta=f"{cnt_a} / {cnt_b}")

    st.divider()
    
    # --- 3. 三列库存列表 ---
    c_chart, c_table = st.columns([1, 1])
    
    with c_chart:
        total_all = len(display_df)
        
        # --- V7.4 Display Breakdown ---
        cnt_instock = len(display_df[display_df['状态'] == '库存中'])
        cnt_pending = len(display_df[display_df['状态'] == '待入库'])
        
        st.metric("📦 当前总库存 (Total)", f"{total_all} 台")
        
        # Sub-metrics
        cm1, cm2 = st.columns(2)
        with cm1: st.metric("✅ 在库 (In Stock)", f"{cnt_instock}", help="已实际入库的现货")
        with cm2: st.metric("⏳ 待入库 (Pending)", f"{cnt_pending}", help="流水号已生成但未入库")
        
        if PLOTLY_AVAILABLE and not display_df.empty and px:
            fig = px.pie(display_df, names='机型', hole=0.4, title="机型分布")
            st.plotly_chart(fig, use_container_width=True)

    with c_table:
        # 计算三列数据
        # Group by Model and Status
        if not display_df.empty:
            stats = display_df.groupby(['机型', '状态']).size().unstack(fill_value=0)
            if '库存中' not in stats.columns: stats['库存中'] = 0
            if '待入库' not in stats.columns: stats['待入库'] = 0
        else:
            stats = pd.DataFrame(columns=['库存中', '待入库'])

        summary_data = []
        for m in display_models:
            in_stock = 0
            pending = 0
            
            if m in stats.index:
                in_stock = int(stats.loc[m, '库存中'])
                pending = int(stats.loc[m, '待入库'])
            
            total = in_stock + pending
            
            # 显示逻辑：有库存 OR 是重点机型
            is_key_model = m in CUSTOM_MODEL_ORDER
            if total > 0 or is_key_model or selected_models_query:
                summary_data.append({
                    "机型": m,
                    "库存中": in_stock,
                    "待入库": pending,
                    "全部": total
                })
        
        summary_df = pd.DataFrame(summary_data)
        if not summary_df.empty:
            summary_df['__rank'] = summary_df['机型'].apply(get_model_rank)
            summary_df = summary_df.sort_values(by=['__rank'], ascending=True)
            
            st.dataframe(
                summary_df.drop(columns=['__rank']), 
                use_container_width=True, 
                hide_index=True,
                column_config={
                    "库存中": st.column_config.NumberColumn(format="%d"),
                    "待入库": st.column_config.NumberColumn(format="%d"),
                    "全部": st.column_config.NumberColumn(format="%d"),
                }
            )
        else:
            st.info("无数据")
            
        with st.expander("详细清单 (Detailed List)"):
            # 防止SettingWithCopyWarning
            display_df = display_df.copy()
            display_df['__rank'] = display_df['机型'].apply(get_model_rank)
            display_df = display_df.sort_values(by=['__rank', '批次号'], ascending=[True, False])
            st.dataframe(display_df.drop(columns=['__rank'])[['批次号', '机型', '流水号', '状态', '机台备注/配置']], use_container_width=True)

# --- 📜 日志 ---
elif st.session_state.page == 'log_viewer':
    c_back, c_title = st.columns([1.5, 8.5])
    with c_back: st.button("⬅️ 返回", on_click=go_home, use_container_width=True)
    with c_title: st.header("📜 日志")
    try:
        with get_engine().connect() as conn:
            df = pd.read_sql("SELECT * FROM transaction_log ", conn)
        st.dataframe(df, use_container_width=True)
    except: st.info("暂无日志")
